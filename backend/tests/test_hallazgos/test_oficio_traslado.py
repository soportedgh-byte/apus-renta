"""
Tests para generacion de oficios de traslado — Sprint 5.
Verifica la generacion DOCX por connotacion (fiscal, disciplinario, penal).
"""

import io
import pytest
from unittest.mock import AsyncMock, MagicMock
from datetime import datetime, timezone

from docx import Document

from app.services.hallazgo_service import (
    HallazgoService,
    DESTINO_TRASLADO,
    _OficioTrasladoDocx,
)


def _mock_hallazgo_aprobado(**kwargs):
    """Crea un hallazgo mock en estado APROBADO."""
    h = MagicMock()
    h.id = "h-traslado-001"
    h.auditoria_id = "aud-001"
    h.numero_hallazgo = 1
    h.titulo = "Sobrecostos en contratacion"
    h.condicion = "Se identificaron sobrecostos del 35% en 12 contratos."
    h.criterio = "Ley 80 de 1993, art. 25, num. 12."
    h.causa = "Deficiencia en estudios de mercado."
    h.efecto = "Presunto detrimento por $2.345.678.900."
    h.recomendacion = "Fortalecer estudios de mercado."
    h.connotaciones = kwargs.get("connotaciones", [
        {"tipo": "fiscal", "fundamentacion_legal": "Ley 610 de 2000"},
    ])
    h.cuantia_presunto_dano = kwargs.get("cuantia_presunto_dano", 2345678900)
    h.presuntos_responsables = kwargs.get("presuntos_responsables", [
        {"nombre": "Director", "cargo": "Director de Contratacion", "periodo": "2024-2025", "fundamentacion": "Ordenador"},
    ])
    h.evidencias = kwargs.get("evidencias", [
        {"documento_id": "d1", "descripcion": "Contratos 001-012", "folio": "1-120", "tipo": "documental"},
    ])
    h.estado = kwargs.get("estado", "APROBADO")
    h.fase_actual_workflow = "director"
    h.historial_workflow = []
    h.redaccion_validada_humano = True
    h.generado_por_ia = True
    h.validado_por = 2
    h.fecha_validacion = datetime.now(timezone.utc)
    h.created_by = 1
    h.updated_by = 2
    h.created_at = datetime.now(timezone.utc)
    h.updated_at = datetime.now(timezone.utc)
    return h


class TestDestinosTraslado:
    """Tests para la configuracion de destinos."""

    def test_3_destinos_configurados(self):
        """Hay exactamente 3 destinos: fiscal, disciplinario, penal."""
        assert len(DESTINO_TRASLADO) == 3

    def test_destino_fiscal(self):
        """El destino fiscal es la Contraloria Delegada."""
        info = DESTINO_TRASLADO["fiscal"]
        assert "Contraloria Delegada" in info["entidad"]
        assert "610" in info["fundamentacion"]

    def test_destino_disciplinario(self):
        """El destino disciplinario es la Procuraduria."""
        info = DESTINO_TRASLADO["disciplinario"]
        assert "Procuraduria" in info["entidad"]
        assert "734" in info["fundamentacion"]

    def test_destino_penal(self):
        """El destino penal es la Fiscalia."""
        info = DESTINO_TRASLADO["penal"]
        assert "Fiscalia" in info["entidad"]


class TestOficioTrasladoDocx:
    """Tests para la generacion del documento DOCX."""

    def test_generar_oficio_fiscal(self):
        """Genera DOCX de oficio fiscal valido."""
        h = _mock_hallazgo_aprobado()
        generador = _OficioTrasladoDocx(h, "fiscal", DESTINO_TRASLADO["fiscal"])
        resultado = generador.generar_bytes()

        assert isinstance(resultado, bytes)
        assert len(resultado) > 0

        doc = Document(io.BytesIO(resultado))
        textos = " ".join(p.text for p in doc.paragraphs).upper()
        assert "TRASLADO" in textos

    def test_generar_oficio_disciplinario(self):
        """Genera DOCX de oficio disciplinario valido."""
        h = _mock_hallazgo_aprobado()
        generador = _OficioTrasladoDocx(h, "disciplinario", DESTINO_TRASLADO["disciplinario"])
        resultado = generador.generar_bytes()

        assert isinstance(resultado, bytes)
        doc = Document(io.BytesIO(resultado))
        assert len(doc.paragraphs) > 0

    def test_generar_oficio_penal(self):
        """Genera DOCX de oficio penal valido."""
        h = _mock_hallazgo_aprobado()
        generador = _OficioTrasladoDocx(h, "penal", DESTINO_TRASLADO["penal"])
        resultado = generador.generar_bytes()

        assert isinstance(resultado, bytes)
        doc = Document(io.BytesIO(resultado))
        assert len(doc.paragraphs) > 0

    def test_oficio_contiene_4_elementos(self):
        """El oficio contiene las 4 secciones obligatorias."""
        h = _mock_hallazgo_aprobado()
        generador = _OficioTrasladoDocx(h, "fiscal", DESTINO_TRASLADO["fiscal"])
        resultado = generador.generar_bytes()

        doc = Document(io.BytesIO(resultado))
        textos = " ".join(p.text for p in doc.paragraphs).upper()
        assert "CONDICION" in textos
        assert "CRITERIO" in textos
        assert "CAUSA" in textos
        assert "EFECTO" in textos

    def test_oficio_contiene_cuantia(self):
        """El oficio incluye la cuantificacion del dano."""
        h = _mock_hallazgo_aprobado(cuantia_presunto_dano=5000000000)
        generador = _OficioTrasladoDocx(h, "fiscal", DESTINO_TRASLADO["fiscal"])
        resultado = generador.generar_bytes()

        doc = Document(io.BytesIO(resultado))
        textos = " ".join(p.text for p in doc.paragraphs).upper()
        # Debe contener referencia a la cuantificacion
        assert "CUANTIFICACION" in textos or "DANO" in textos

    def test_oficio_tiene_tablas(self):
        """El oficio debe tener multiples tablas."""
        h = _mock_hallazgo_aprobado()
        generador = _OficioTrasladoDocx(h, "fiscal", DESTINO_TRASLADO["fiscal"])
        resultado = generador.generar_bytes()

        doc = Document(io.BytesIO(resultado))
        # Debe tener al menos 3 tablas: encabezado, datos, responsables/evidencias
        assert len(doc.tables) >= 3

    def test_oficio_sin_cuantia(self):
        """Genera oficio sin cuantia (hallazgo administrativo convertido)."""
        h = _mock_hallazgo_aprobado(cuantia_presunto_dano=None)
        generador = _OficioTrasladoDocx(h, "disciplinario", DESTINO_TRASLADO["disciplinario"])
        resultado = generador.generar_bytes()

        assert isinstance(resultado, bytes)
        assert len(resultado) > 0

    def test_oficio_sin_responsables(self):
        """Genera oficio sin presuntos responsables."""
        h = _mock_hallazgo_aprobado(presuntos_responsables=[])
        generador = _OficioTrasladoDocx(h, "fiscal", DESTINO_TRASLADO["fiscal"])
        resultado = generador.generar_bytes()

        assert isinstance(resultado, bytes)
        assert len(resultado) > 0


class TestValidacionEstadoTraslado:
    """Tests para la validacion de estado al generar oficio."""

    @pytest.mark.asyncio
    async def test_oficio_solo_para_aprobados(self):
        """Solo se puede generar oficio para hallazgos APROBADOS."""
        db = AsyncMock()
        h = _mock_hallazgo_aprobado(estado="EN_REVISION")

        result_mock = MagicMock()
        result_mock.scalar_one_or_none.return_value = h
        db.execute.return_value = result_mock

        servicio = HallazgoService(db)
        with pytest.raises(ValueError, match="APROBADOS"):
            await servicio.generar_oficio_traslado("h-001", "fiscal")

    @pytest.mark.asyncio
    async def test_destino_invalido(self):
        """Destino invalido lanza error."""
        db = AsyncMock()
        h = _mock_hallazgo_aprobado(estado="APROBADO")

        result_mock = MagicMock()
        result_mock.scalar_one_or_none.return_value = h
        db.execute.return_value = result_mock

        servicio = HallazgoService(db)
        with pytest.raises(ValueError, match="no valido"):
            await servicio.generar_oficio_traslado("h-001", "invalido")
