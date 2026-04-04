"""
Tests para Formato F14 — Matriz de Riesgos de Auditoria.
Sprint: 4
"""

import io
import pytest
from docx import Document

from app.formatos.formato_14_matriz_riesgos import FormatoF14MatrizRiesgos


class TestFormatoF14:
    """Tests para el formato F14 — Matriz de Riesgos."""

    def test_generar_vacio(self):
        """Genera F14 sin datos — todos los campos [COMPLETAR]."""
        fmt = FormatoF14MatrizRiesgos()
        resultado = fmt.generar_bytes()
        assert isinstance(resultado, bytes)
        assert len(resultado) > 0
        doc = Document(io.BytesIO(resultado))
        textos = " ".join(p.text for p in doc.paragraphs)
        assert "MATRIZ DE RIESGOS" in textos
        assert "CONTRALORIA" in textos

    def test_generar_con_datos(self):
        """Genera F14 con datos de riesgos prellenados."""
        datos = {
            "nombre_entidad": "MinTIC",
            "vigencia": "2025",
            "tipo_auditoria": "Financiera y de Gestion",
            "riesgos": [
                {
                    "id": "R01",
                    "proceso": "Presupuesto",
                    "descripcion": "Sobrestimacion de ingresos",
                    "tipo": "Inherente",
                    "probabilidad": "4",
                    "impacto": "4",
                    "riesgo_inherente": "Alto",
                    "control_existente": "Revision mensual",
                    "efectividad_control": "Parcial",
                    "riesgo_residual": "Medio",
                    "nivel_riesgo": "Alto",
                    "respuesta_auditor": "Prueba sustantiva de ingresos",
                },
                {
                    "id": "R02",
                    "proceso": "Contratacion",
                    "descripcion": "Fraccionamiento de contratos",
                    "tipo": "Control",
                    "probabilidad": "3",
                    "impacto": "5",
                    "riesgo_inherente": "Alto",
                    "control_existente": "Comite de contratacion",
                    "efectividad_control": "Efectivo",
                    "riesgo_residual": "Bajo",
                    "nivel_riesgo": "Medio",
                    "respuesta_auditor": "Revision de expedientes",
                },
            ],
        }
        fmt = FormatoF14MatrizRiesgos(datos=datos)
        resultado = fmt.generar_bytes()
        doc = Document(io.BytesIO(resultado))

        textos = " ".join(p.text for p in doc.paragraphs)
        assert "MinTIC" in textos or "MINTIC" in textos.upper()

        # Verificar que tiene tablas (escalas + matriz + mapa calor + firmas)
        assert len(doc.tables) >= 4

    def test_tiene_mapa_calor(self):
        """Verifica que el mapa de calor de riesgos esta presente."""
        fmt = FormatoF14MatrizRiesgos(datos={"nombre_entidad": "Test"})
        resultado = fmt.generar_bytes()
        doc = Document(io.BytesIO(resultado))

        textos = " ".join(p.text for p in doc.paragraphs)
        assert "MAPA DE CALOR" in textos

    def test_tiene_escalas_valoracion(self):
        """Verifica las escalas de probabilidad e impacto."""
        fmt = FormatoF14MatrizRiesgos()
        resultado = fmt.generar_bytes()
        doc = Document(io.BytesIO(resultado))

        textos = " ".join(p.text for p in doc.paragraphs)
        assert "PROBABILIDAD" in textos
        assert "IMPACTO" in textos

    def test_tiene_seccion_firmas(self):
        """Verifica que la seccion de firmas existe."""
        fmt = FormatoF14MatrizRiesgos()
        resultado = fmt.generar_bytes()
        doc = Document(io.BytesIO(resultado))

        textos = " ".join(p.text for p in doc.paragraphs)
        assert "FIRMAS" in textos
