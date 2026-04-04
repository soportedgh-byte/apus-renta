"""
Tests para Formato F17 — Calculo de Materialidad.
Sprint: 4
"""

import io
import pytest
from docx import Document

from app.formatos.formato_17_materialidad import FormatoF17Materialidad


class TestFormatoF17:
    """Tests para el formato F17 — Calculo de Materialidad."""

    def test_generar_vacio(self):
        """Genera F17 sin datos."""
        fmt = FormatoF17Materialidad()
        resultado = fmt.generar_bytes()
        assert isinstance(resultado, bytes)
        assert len(resultado) > 0

    def test_generar_con_datos_completos(self):
        """Genera F17 con datos de materialidad prellenados."""
        datos = {
            "nombre_entidad": "MinTIC",
            "vigencia": "2025",
            "tipo_auditoria": "Financiera",
            "total_activos": "500,000,000,000",
            "total_ingresos": "300,000,000,000",
            "total_gastos": "280,000,000,000",
            "patrimonio": "220,000,000,000",
            "presupuesto": "350,000,000,000",
            "base_seleccionada": "Total Activos",
            "valor_base": "500,000,000,000",
            "porcentaje_aplicado": "1.5%",
            "materialidad_global": "7,500,000,000",
            "pct_ejecucion": "60%",
            "materialidad_ejecucion": "4,500,000,000",
            "pct_umbral": "5%",
            "umbral_insignificante": "375,000,000",
        }
        fmt = FormatoF17Materialidad(datos=datos)
        resultado = fmt.generar_bytes()
        doc = Document(io.BytesIO(resultado))

        textos = " ".join(p.text for p in doc.paragraphs)
        assert "MATERIALIDAD" in textos
        assert "NIA 320" in textos

    def test_contiene_nia_320_450(self):
        """Verifica referencia a NIA 320 y NIA 450."""
        fmt = FormatoF17Materialidad()
        resultado = fmt.generar_bytes()
        doc = Document(io.BytesIO(resultado))

        textos = " ".join(p.text for p in doc.paragraphs)
        assert "NIA 320" in textos
        assert "NIA 450" in textos

    def test_tiene_tabla_bases_calculo(self):
        """Verifica que existe la tabla de bases de calculo."""
        fmt = FormatoF17Materialidad()
        resultado = fmt.generar_bytes()
        doc = Document(io.BytesIO(resultado))

        # Debe tener varias tablas: encabezado, bases, calculo, ejecucion, umbral, resumen, firmas
        assert len(doc.tables) >= 5

    def test_tiene_resumen_niveles(self):
        """Verifica la tabla resumen de niveles de materialidad."""
        fmt = FormatoF17Materialidad()
        resultado = fmt.generar_bytes()
        doc = Document(io.BytesIO(resultado))

        textos = " ".join(p.text for p in doc.paragraphs)
        assert "RESUMEN" in textos

    def test_formato_es_f17(self):
        """Verifica el numero correcto del formato."""
        fmt = FormatoF17Materialidad()
        assert fmt.numero_formato == 17
        resultado = fmt.generar_bytes()
        doc = Document(io.BytesIO(resultado))

        textos = " ".join(p.text for p in doc.paragraphs)
        assert "FORMATO F17" in textos
