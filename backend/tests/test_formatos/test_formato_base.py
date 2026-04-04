"""
Tests para FormatoBaseCGR — Clase base de generacion DOCX.
Sprint: 4
"""

import io
import pytest
from docx import Document

from app.formatos.formato_base import FormatoBaseCGR


class FormatoTestCGR(FormatoBaseCGR):
    """Formato de prueba que implementa construir()."""

    def construir(self) -> None:
        self.agregar_titulo_seccion("Seccion de Prueba")
        self.agregar_parrafo("Parrafo de prueba con contenido.")
        self.agregar_parrafo_justificado("Parrafo justificado de prueba.")
        self.agregar_campo_completar("campo de prueba")
        self.crear_tabla(
            encabezados=["Col A", "Col B", "Col C"],
            filas=[
                ["Dato 1", "Dato 2", "Dato 3"],
                ["[COMPLETAR]", "Dato 5", "[COMPLETAR]"],
            ],
        )
        self.crear_tabla_clave_valor([
            ("Clave 1", "Valor 1"),
            ("Clave 2", "[COMPLETAR]"),
        ])
        self.agregar_seccion_firmas()


class TestFormatoBaseCGR:
    """Tests para la clase base FormatoBaseCGR."""

    def test_instanciacion(self):
        fmt = FormatoTestCGR(
            numero_formato=99,
            nombre_formato="Test",
            datos={"nombre_entidad": "Entidad Test", "vigencia": "2025"},
        )
        assert fmt.numero_formato == 99
        assert fmt.nombre_formato == "Test"
        assert fmt.datos["nombre_entidad"] == "Entidad Test"

    def test_generar_bytes_retorna_bytes(self):
        fmt = FormatoTestCGR(
            numero_formato=99,
            nombre_formato="Test",
            datos={"nombre_entidad": "MinTIC", "vigencia": "2025"},
        )
        resultado = fmt.generar_bytes()
        assert isinstance(resultado, bytes)
        assert len(resultado) > 0

    def test_generar_bytes_es_docx_valido(self):
        fmt = FormatoTestCGR(
            numero_formato=99,
            nombre_formato="Test",
            datos={"nombre_entidad": "MinTIC", "vigencia": "2025"},
        )
        resultado = fmt.generar_bytes()
        # Verificar que es un DOCX valido (ZIP con magic number PK)
        assert resultado[:2] == b"PK"
        # Verificar que se puede abrir con python-docx
        doc = Document(io.BytesIO(resultado))
        assert len(doc.paragraphs) > 0

    def test_encabezado_contiene_cgr(self):
        fmt = FormatoTestCGR(
            numero_formato=99,
            nombre_formato="Test",
            datos={"nombre_entidad": "MinTIC", "vigencia": "2025"},
        )
        resultado = fmt.generar_bytes()
        doc = Document(io.BytesIO(resultado))

        textos = [p.text for p in doc.paragraphs]
        assert any("CONTRALORIA" in t for t in textos)

    def test_encabezado_contiene_formato(self):
        fmt = FormatoTestCGR(
            numero_formato=99,
            nombre_formato="Test",
            datos={},
        )
        resultado = fmt.generar_bytes()
        doc = Document(io.BytesIO(resultado))

        textos = [p.text for p in doc.paragraphs]
        assert any("FORMATO F99" in t for t in textos)

    def test_campo_por_completar(self):
        fmt = FormatoTestCGR(numero_formato=1, nombre_formato="Test", datos={})
        assert fmt.campo_por_completar() == "[COMPLETAR]"
        assert fmt.campo_por_completar("dato") == "[COMPLETAR — dato]"

    def test_tiene_tablas(self):
        fmt = FormatoTestCGR(
            numero_formato=99,
            nombre_formato="Test",
            datos={"nombre_entidad": "Test"},
        )
        resultado = fmt.generar_bytes()
        doc = Document(io.BytesIO(resultado))
        # Debe tener al menos: tabla de datos auditoria, tabla contenido,
        # tabla clave-valor, tabla firmas
        assert len(doc.tables) >= 3

    def test_pie_pagina_existe(self):
        fmt = FormatoTestCGR(
            numero_formato=99,
            nombre_formato="Test",
            datos={},
        )
        resultado = fmt.generar_bytes()
        doc = Document(io.BytesIO(resultado))

        # Verificar que el footer tiene contenido
        for seccion in doc.sections:
            footer = seccion.footer
            textos_footer = [p.text for p in footer.paragraphs]
            texto_completo = " ".join(textos_footer)
            assert "Circular 023" in texto_completo or "Formato F99" in texto_completo

    def test_nota_ia_presente(self):
        fmt = FormatoTestCGR(
            numero_formato=99,
            nombre_formato="Test",
            datos={},
        )
        resultado = fmt.generar_bytes()
        doc = Document(io.BytesIO(resultado))

        textos = [p.text for p in doc.paragraphs]
        assert any("inteligencia artificial" in t.lower() for t in textos)

    def test_construir_no_implementado_lanza_error(self):
        fmt = FormatoBaseCGR(
            numero_formato=99,
            nombre_formato="Base",
            datos={},
        )
        with pytest.raises(NotImplementedError):
            fmt.construir()
