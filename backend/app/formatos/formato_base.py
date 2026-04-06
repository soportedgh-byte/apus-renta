"""
CecilIA v2 — Sistema de IA para Control Fiscal
Contraloria General de la Republica de Colombia

Archivo: app/formatos/formato_base.py
Proposito: Clase base FormatoBaseCGR — generacion de documentos DOCX
           con encabezado institucional CGR, estilos profesionales,
           pie de pagina con disclaimer y campos por completar
Sprint: 4
Autor: Equipo Tecnico CecilIA — CD-TIC-CGR
Fecha: Abril 2026
"""

from __future__ import annotations

import io
import logging
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Inches, Mm, Pt, RGBColor

logger = logging.getLogger("cecilia.formatos.base")

# ── Constantes de estilo ────────────────────────────────────────────────────

COLOR_AZUL_CGR = RGBColor(0x1A, 0x52, 0x76)       # #1A5276 — headers tablas
COLOR_AZUL_CLARO = RGBColor(0xEB, 0xF5, 0xFB)     # #EBF5FB — filas alternas
COLOR_ORO = RGBColor(0xC9, 0xA8, 0x4C)            # #C9A84C — linea dorada
COLOR_ROJO = RGBColor(0xC0, 0x39, 0x2B)            # #C0392B — campos pendientes
COLOR_BLANCO = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_NEGRO = RGBColor(0x00, 0x00, 0x00)
COLOR_GRIS = RGBColor(0x5F, 0x63, 0x68)            # #5F6368 — texto secundario
COLOR_GRIS_CLARO = RGBColor(0x95, 0x95, 0x95)

FUENTE_CUERPO = "Calibri"
FUENTE_TITULO = "Georgia"
TAMANO_CUERPO = Pt(11)
TAMANO_TITULO_PRINCIPAL = Pt(14)
TAMANO_TITULO_SECCION = Pt(12)
TAMANO_PEQUENO = Pt(9)
TAMANO_MUY_PEQUENO = Pt(8)


class FormatoBaseCGR:
    """Clase base para generacion de formatos CGR en DOCX profesional.

    Provee metodos comunes para:
    - Encabezado institucional con logo CGR
    - Linea separadora dorada
    - Tabla de datos de auditoria
    - Estilos de tablas profesionales
    - Campos por completar en rojo
    - Pie de pagina con numeracion y disclaimer IA
    - Generacion de bytes del documento
    """

    def __init__(
        self,
        numero_formato: int,
        nombre_formato: str,
        datos: dict[str, Any] | None = None,
    ) -> None:
        self.numero_formato = numero_formato
        self.nombre_formato = nombre_formato
        self.datos = datos or {}
        self.doc = Document()
        self._configurar_estilos()
        self._configurar_margenes()

    # ── Configuracion inicial ───────────────────────────────────────────────

    def _configurar_margenes(self) -> None:
        """Margenes estandar: 2.5cm superior/inferior, 2cm laterales."""
        for seccion in self.doc.sections:
            seccion.top_margin = Cm(2.5)
            seccion.bottom_margin = Cm(2.5)
            seccion.left_margin = Cm(2.0)
            seccion.right_margin = Cm(2.0)
            seccion.page_width = Mm(210)   # A4
            seccion.page_height = Mm(297)

    def _configurar_estilos(self) -> None:
        """Configura estilos base del documento."""
        estilo = self.doc.styles["Normal"]
        fuente = estilo.font
        fuente.name = FUENTE_CUERPO
        fuente.size = TAMANO_CUERPO
        fuente.color.rgb = COLOR_NEGRO

        formato = estilo.paragraph_format
        formato.space_before = Pt(2)
        formato.space_after = Pt(4)
        formato.line_spacing = 1.15

    # ── Encabezado institucional ────────────────────────────────────────────

    def agregar_encabezado(
        self,
        entidad: str | None = None,
        vigencia: str | None = None,
        tipo_auditoria: str | None = None,
        auditoria_id: str | None = None,
    ) -> None:
        """Agrega encabezado institucional CGR completo.

        Incluye:
        1. Logo CGR (si existe)
        2. Titulo "CONTRALORIA GENERAL DE LA REPUBLICA"
        3. Subtitulo "Sistema de Control Fiscal"
        4. Linea separadora dorada (#C9A84C)
        5. Titulo del formato (Georgia bold 14pt)
        6. Tabla de datos de auditoria
        """
        entidad = entidad or self.datos.get("nombre_entidad", "[COMPLETAR]")
        vigencia = vigencia or self.datos.get("vigencia", "[COMPLETAR]")
        tipo_auditoria = tipo_auditoria or self.datos.get("tipo_auditoria", "")
        auditoria_id = auditoria_id or self.datos.get("auditoria_id", "")

        # Logo CGR
        self._agregar_logo()

        # Titulo institucional
        p_titulo = self.doc.add_paragraph()
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_titulo.add_run("CONTRALORIA GENERAL DE LA REPUBLICA")
        run.font.name = FUENTE_TITULO
        run.font.size = TAMANO_TITULO_PRINCIPAL
        run.font.bold = True
        run.font.color.rgb = COLOR_AZUL_CGR

        # Subtitulo
        p_sub = self.doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.paragraph_format.space_before = Pt(0)
        p_sub.paragraph_format.space_after = Pt(6)
        run_sub = p_sub.add_run("Sistema de Control Fiscal — CecilIA v2")
        run_sub.font.name = FUENTE_CUERPO
        run_sub.font.size = TAMANO_PEQUENO
        run_sub.font.color.rgb = COLOR_GRIS

        # Linea dorada separadora
        self._agregar_linea_dorada()

        # Titulo del formato
        p_formato = self.doc.add_paragraph()
        p_formato.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_formato.paragraph_format.space_before = Pt(10)
        p_formato.paragraph_format.space_after = Pt(6)
        run_fmt = p_formato.add_run(
            f"FORMATO F{self.numero_formato:02d} — {self.nombre_formato.upper()}"
        )
        run_fmt.font.name = FUENTE_TITULO
        run_fmt.font.size = TAMANO_TITULO_SECCION
        run_fmt.font.bold = True
        run_fmt.font.color.rgb = COLOR_AZUL_CGR

        # Tabla de datos de auditoria
        fecha_actual = datetime.now().strftime("%d/%m/%Y")
        datos_auditoria = [
            ("Entidad auditada", entidad),
            ("Vigencia auditada", vigencia),
        ]
        if tipo_auditoria:
            datos_auditoria.append(("Tipo de auditoria", tipo_auditoria))
        if auditoria_id:
            datos_auditoria.append(("ID Auditoria", auditoria_id))
        datos_auditoria.append(("Fecha de elaboracion", fecha_actual))

        tabla = self.doc.add_table(rows=len(datos_auditoria), cols=2)
        tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._estilo_tabla_datos(tabla)

        for i, (campo, valor) in enumerate(datos_auditoria):
            celda_campo = tabla.rows[i].cells[0]
            celda_valor = tabla.rows[i].cells[1]
            self._escribir_celda(celda_campo, campo, negrita=True, tamano=TAMANO_PEQUENO)
            if valor == "[COMPLETAR]":
                self._escribir_campo_completar(celda_valor)
            else:
                self._escribir_celda(celda_valor, valor, tamano=TAMANO_PEQUENO)

        # Espaciado despues del encabezado
        self.doc.add_paragraph()

    def _agregar_logo(self) -> None:
        """Intenta agregar el logo CGR al documento."""
        rutas_logo = [
            Path(__file__).parent.parent / "static" / "logo-cgr.png",
            Path(__file__).parent.parent.parent / "static" / "logo-cgr.png",
            Path(__file__).parent.parent / "assets" / "logo-cgr.png",
            Path("/app/app/static/logo-cgr.png"),
        ]
        for ruta in rutas_logo:
            if ruta.exists():
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run()
                run.add_picture(str(ruta), width=Cm(8))
                return

        # Si no hay logo, agregar texto placeholder
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run("[ ESCUDO CGR ]")
        run.font.size = TAMANO_PEQUENO
        run.font.color.rgb = COLOR_GRIS

    def _agregar_linea_dorada(self) -> None:
        """Agrega una linea horizontal dorada (#C9A84C) como separador."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

        # Crear linea con borde inferior del parrafo
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")  # 1.5pt
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "C9A84C")
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ── Pie de pagina ───────────────────────────────────────────────────────

    def agregar_pie_pagina(self) -> None:
        """Agrega pie de pagina con numeracion y disclaimer de IA.

        Formato:
        Izquierda: "Formato F{XX} — {nombre} | Generado por CecilIA v2"
        Centro: Pagina {n} de {total}
        Derecha: "Documento asistido por IA — Requiere validacion humana (Circular 023)"
        """
        for seccion in self.doc.sections:
            footer = seccion.footer
            footer.is_linked_to_previous = False

            # Parrafo principal del pie
            if footer.paragraphs:
                p = footer.paragraphs[0]
            else:
                p = footer.add_paragraph()

            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)

            # Linea superior del footer
            pPr = p._element.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            top = OxmlElement("w:top")
            top.set(qn("w:val"), "single")
            top.set(qn("w:sz"), "6")
            top.set(qn("w:space"), "4")
            top.set(qn("w:color"), "C9A84C")
            pBdr.append(top)
            pPr.append(pBdr)

            # Texto formato
            run1 = p.add_run(
                f"Formato F{self.numero_formato:02d} — {self.nombre_formato}"
            )
            run1.font.name = FUENTE_CUERPO
            run1.font.size = TAMANO_MUY_PEQUENO
            run1.font.color.rgb = COLOR_GRIS

            run_sep = p.add_run("   |   ")
            run_sep.font.size = TAMANO_MUY_PEQUENO
            run_sep.font.color.rgb = COLOR_GRIS_CLARO

            # Numeracion de pagina
            run_pag = p.add_run("Pagina ")
            run_pag.font.name = FUENTE_CUERPO
            run_pag.font.size = TAMANO_MUY_PEQUENO
            run_pag.font.color.rgb = COLOR_GRIS

            # Campo de pagina actual
            fldChar1 = OxmlElement("w:fldChar")
            fldChar1.set(qn("w:fldCharType"), "begin")
            run_num = p.add_run()
            run_num._element.append(fldChar1)

            instrText = OxmlElement("w:instrText")
            instrText.set(qn("xml:space"), "preserve")
            instrText.text = " PAGE "
            run_num2 = p.add_run()
            run_num2._element.append(instrText)

            fldChar2 = OxmlElement("w:fldChar")
            fldChar2.set(qn("w:fldCharType"), "end")
            run_num3 = p.add_run()
            run_num3._element.append(fldChar2)

            run_de = p.add_run(" de ")
            run_de.font.name = FUENTE_CUERPO
            run_de.font.size = TAMANO_MUY_PEQUENO
            run_de.font.color.rgb = COLOR_GRIS

            # Campo de total paginas
            fldChar3 = OxmlElement("w:fldChar")
            fldChar3.set(qn("w:fldCharType"), "begin")
            run_tot = p.add_run()
            run_tot._element.append(fldChar3)

            instrText2 = OxmlElement("w:instrText")
            instrText2.set(qn("xml:space"), "preserve")
            instrText2.text = " NUMPAGES "
            run_tot2 = p.add_run()
            run_tot2._element.append(instrText2)

            fldChar4 = OxmlElement("w:fldChar")
            fldChar4.set(qn("w:fldCharType"), "end")
            run_tot3 = p.add_run()
            run_tot3._element.append(fldChar4)

            # Disclaimer
            p_disc = footer.add_paragraph()
            p_disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_disc = p_disc.add_run(
                "Documento asistido por IA — Requiere validacion humana (Circular 023 CGR)"
            )
            run_disc.font.name = FUENTE_CUERPO
            run_disc.font.size = Pt(7)
            run_disc.font.color.rgb = COLOR_GRIS_CLARO
            run_disc.font.italic = True

    # ── Metodos de contenido ────────────────────────────────────────────────

    def agregar_titulo_seccion(self, titulo: str) -> None:
        """Agrega un titulo de seccion con estilo Georgia 12pt bold azul."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(titulo.upper())
        run.font.name = FUENTE_TITULO
        run.font.size = TAMANO_TITULO_SECCION
        run.font.bold = True
        run.font.color.rgb = COLOR_AZUL_CGR

        # Linea inferior sutil
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), "1A5276")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def agregar_subtitulo(self, texto: str) -> None:
        """Agrega un subtitulo Calibri 11pt bold."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(texto)
        run.font.name = FUENTE_CUERPO
        run.font.size = TAMANO_CUERPO
        run.font.bold = True
        run.font.color.rgb = COLOR_AZUL_CGR

    def agregar_parrafo(self, texto: str, negrita: bool = False) -> None:
        """Agrega un parrafo de texto Calibri 11pt."""
        p = self.doc.add_paragraph()
        run = p.add_run(texto)
        run.font.name = FUENTE_CUERPO
        run.font.size = TAMANO_CUERPO
        run.font.bold = negrita

    def agregar_parrafo_justificado(self, texto: str) -> None:
        """Agrega un parrafo justificado."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(texto)
        run.font.name = FUENTE_CUERPO
        run.font.size = TAMANO_CUERPO

    def campo_por_completar(self, descripcion: str = "") -> str:
        """Retorna texto [COMPLETAR] para campos pendientes.

        Se renderiza en rojo bold (#C0392B) cuando se usa con
        _escribir_campo_completar().
        """
        if descripcion:
            return f"[COMPLETAR — {descripcion}]"
        return "[COMPLETAR]"

    def agregar_campo_completar(self, descripcion: str = "") -> None:
        """Agrega un campo [COMPLETAR] en rojo bold al documento."""
        p = self.doc.add_paragraph()
        texto = self.campo_por_completar(descripcion)
        run = p.add_run(texto)
        run.font.name = FUENTE_CUERPO
        run.font.size = TAMANO_CUERPO
        run.font.bold = True
        run.font.color.rgb = COLOR_ROJO

    # ── Tablas profesionales ────────────────────────────────────────────────

    def crear_tabla(
        self,
        encabezados: list[str],
        filas: list[list[str]],
        anchos: list[float] | None = None,
    ) -> Any:
        """Crea una tabla profesional con encabezados azul CGR y filas alternas.

        Args:
            encabezados: Lista de textos para la fila de encabezado.
            filas: Lista de listas con datos de cada fila.
            anchos: Anchos en cm para cada columna (opcional).

        Returns:
            Objeto tabla de python-docx.
        """
        tabla = self.doc.add_table(rows=1 + len(filas), cols=len(encabezados))
        tabla.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Encabezados
        for j, enc in enumerate(encabezados):
            celda = tabla.rows[0].cells[j]
            self._estilo_celda_encabezado(celda, enc)

        # Filas de datos
        for i, fila in enumerate(filas):
            for j, valor in enumerate(fila):
                celda = tabla.rows[i + 1].cells[j]
                es_alterna = i % 2 == 0
                if valor == "[COMPLETAR]" or valor.startswith("[COMPLETAR"):
                    self._escribir_campo_completar(celda)
                else:
                    self._escribir_celda(
                        celda, valor, tamano=TAMANO_PEQUENO,
                        fondo=COLOR_AZUL_CLARO if es_alterna else None,
                    )
                if es_alterna and not (valor == "[COMPLETAR]" or valor.startswith("[COMPLETAR")):
                    self._establecer_fondo_celda(celda, "EBF5FB")

        # Anchos de columna
        if anchos:
            for j, ancho in enumerate(anchos):
                for row in tabla.rows:
                    row.cells[j].width = Cm(ancho)

        # Bordes
        self._establecer_bordes_tabla(tabla)

        return tabla

    def crear_tabla_clave_valor(
        self,
        pares: list[tuple[str, str]],
        ancho_clave: float = 5.0,
        ancho_valor: float = 12.0,
    ) -> Any:
        """Crea una tabla de dos columnas clave-valor.

        Args:
            pares: Lista de tuplas (clave, valor).
            ancho_clave: Ancho de la columna clave en cm.
            ancho_valor: Ancho de la columna valor en cm.
        """
        tabla = self.doc.add_table(rows=len(pares), cols=2)
        tabla.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, (clave, valor) in enumerate(pares):
            celda_c = tabla.rows[i].cells[0]
            celda_v = tabla.rows[i].cells[1]

            self._escribir_celda(celda_c, clave, negrita=True, tamano=TAMANO_PEQUENO)
            self._establecer_fondo_celda(celda_c, "1A5276")
            # Texto blanco para clave
            for p in celda_c.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = COLOR_BLANCO

            if valor == "[COMPLETAR]" or valor.startswith("[COMPLETAR"):
                self._escribir_campo_completar(celda_v)
            else:
                self._escribir_celda(celda_v, valor, tamano=TAMANO_PEQUENO)

            celda_c.width = Cm(ancho_clave)
            celda_v.width = Cm(ancho_valor)

        self._establecer_bordes_tabla(tabla)
        return tabla

    def agregar_seccion_firmas(self) -> None:
        """Agrega seccion de firmas estandar: Elaboro, Reviso, Aprobo."""
        self.agregar_titulo_seccion("Firmas")

        firmas = [
            ("Elaboro", "[COMPLETAR — Nombre del auditor]", "[COMPLETAR]"),
            ("Reviso", "[COMPLETAR — Supervisor]", "[COMPLETAR]"),
            ("Aprobo", "[COMPLETAR — Director Tecnico]", "[COMPLETAR]"),
        ]

        tabla = self.doc.add_table(rows=len(firmas), cols=4)
        tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
        encabezados_firma = ["Rol", "Nombre", "Firma", "Fecha"]

        # Agregar encabezado invisible
        for i, (rol, nombre, fecha) in enumerate(firmas):
            tabla.rows[i].cells[0].text = ""
            self._escribir_celda(tabla.rows[i].cells[0], rol, negrita=True, tamano=TAMANO_PEQUENO)
            self._escribir_campo_completar(tabla.rows[i].cells[1])
            celda_firma = tabla.rows[i].cells[2]
            self._escribir_celda(celda_firma, "________________________", tamano=TAMANO_PEQUENO)
            self._escribir_campo_completar(tabla.rows[i].cells[3])

        self._establecer_bordes_tabla(tabla)

    def agregar_nota_ia(self) -> None:
        """Agrega nota de advertencia sobre contenido generado por IA."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(4)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run(
            "⚠ NOTA: Este documento fue generado con asistencia de inteligencia artificial "
            "(CecilIA v2). Todos los datos, analisis y conclusiones deben ser verificados "
            "y validados por el equipo auditor antes de su uso oficial, conforme a lo "
            "establecido en la Circular 023 de la CGR."
        )
        run.font.name = FUENTE_CUERPO
        run.font.size = TAMANO_PEQUENO
        run.font.italic = True
        run.font.color.rgb = COLOR_GRIS

    # ── Metodos de estilo para celdas ───────────────────────────────────────

    def _estilo_celda_encabezado(self, celda: Any, texto: str) -> None:
        """Aplica estilo de encabezado: fondo azul CGR, texto blanco bold."""
        celda.text = ""
        p = celda.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(texto)
        run.font.name = FUENTE_CUERPO
        run.font.size = TAMANO_PEQUENO
        run.font.bold = True
        run.font.color.rgb = COLOR_BLANCO

        self._establecer_fondo_celda(celda, "1A5276")

    def _estilo_tabla_datos(self, tabla: Any) -> None:
        """Aplica estilo a tabla de datos clave-valor del encabezado."""
        for row in tabla.rows:
            celda_campo = row.cells[0]
            self._establecer_fondo_celda(celda_campo, "1A5276")
            for p in celda_campo.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = COLOR_BLANCO
            celda_campo.width = Cm(5)
            row.cells[1].width = Cm(12)

        self._establecer_bordes_tabla(tabla)

    def _escribir_celda(
        self,
        celda: Any,
        texto: str,
        negrita: bool = False,
        tamano: Any = None,
        fondo: Any = None,
    ) -> None:
        """Escribe texto en una celda con formato."""
        celda.text = ""
        p = celda.paragraphs[0]
        run = p.add_run(texto)
        run.font.name = FUENTE_CUERPO
        run.font.size = tamano or TAMANO_CUERPO
        run.font.bold = negrita

    def _escribir_campo_completar(self, celda: Any, descripcion: str = "") -> None:
        """Escribe [COMPLETAR] en rojo bold en una celda."""
        celda.text = ""
        p = celda.paragraphs[0]
        texto = self.campo_por_completar(descripcion)
        run = p.add_run(texto)
        run.font.name = FUENTE_CUERPO
        run.font.size = TAMANO_PEQUENO
        run.font.bold = True
        run.font.color.rgb = COLOR_ROJO

    def _establecer_fondo_celda(self, celda: Any, color_hex: str) -> None:
        """Establece el color de fondo de una celda."""
        tc = celda._element
        tcPr = tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), color_hex)
        tcPr.append(shading)

    def _establecer_bordes_tabla(self, tabla: Any) -> None:
        """Establece bordes finos para toda la tabla."""
        tbl = tabla._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")

        borders = OxmlElement("w:tblBorders")
        for borde in ("top", "left", "bottom", "right", "insideH", "insideV"):
            elem = OxmlElement(f"w:{borde}")
            elem.set(qn("w:val"), "single")
            elem.set(qn("w:sz"), "4")
            elem.set(qn("w:space"), "0")
            elem.set(qn("w:color"), "CCCCCC")
            borders.append(elem)

        tblPr.append(borders)

    # ── Generacion ──────────────────────────────────────────────────────────

    def construir(self) -> None:
        """Metodo a sobrescribir en cada formato especifico.

        Las subclases deben implementar este metodo para agregar
        el contenido especifico del formato.
        """
        raise NotImplementedError(
            f"El formato F{self.numero_formato:02d} debe implementar construir()"
        )

    def generar(self) -> Document:
        """Genera el documento completo: encabezado + contenido + pie.

        Returns:
            Documento python-docx listo para guardar.
        """
        self.agregar_encabezado(
            entidad=self.datos.get("nombre_entidad"),
            vigencia=self.datos.get("vigencia"),
            tipo_auditoria=self.datos.get("tipo_auditoria"),
            auditoria_id=self.datos.get("auditoria_id"),
        )
        self.construir()
        self.agregar_nota_ia()
        self.agregar_pie_pagina()
        return self.doc

    def generar_bytes(self) -> bytes:
        """Genera el documento y retorna como bytes para streaming.

        Returns:
            Bytes del archivo DOCX listo para descarga.
        """
        self.generar()
        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def guardar(self, ruta: str | Path) -> Path:
        """Genera y guarda el documento en la ruta especificada.

        Args:
            ruta: Ruta donde guardar el archivo DOCX.

        Returns:
            Path del archivo guardado.
        """
        self.generar()
        ruta = Path(ruta)
        ruta.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(ruta))
        logger.info("Formato F%02d guardado en %s", self.numero_formato, ruta)
        return ruta
