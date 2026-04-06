"""
CecilIA v2 — Sistema de IA para Control Fiscal
Contraloria General de la Republica de Colombia

Archivo: app/services/capacitacion_service.py
Proposito: Servicio de gestion de capacitacion: rutas de aprendizaje,
           lecciones, progreso, quizzes y generacion de contenido didactico.
Sprint: 6
Autor: Equipo Tecnico CecilIA — CD-TIC-CGR
Fecha: Abril 2026
"""

from __future__ import annotations

import io
import logging
import uuid
from datetime import datetime, timezone
from typing import Any, Optional

from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.formatos.formato_base import FormatoBaseCGR
from app.models.capacitacion import (
    Leccion,
    ProgresoUsuario,
    QuizResultado,
    RutaAprendizaje,
)

logger = logging.getLogger("cecilia.services.capacitacion")

# ── Seed Data: Rutas y Lecciones ────────────────────────────────────────────

RUTAS_SEED: list[dict[str, Any]] = [
    {
        "id": "ruta-001",
        "nombre": "Conoce la CGR",
        "descripcion": "Aprende sobre la estructura, funciones y macroprocesos de la Contraloria General de la Republica.",
        "icono": "🏛️",
        "color": "#C9A84C",
        "direccion": "TODOS",
        "orden": 1,
        "lecciones": [
            {"num": 1, "titulo": "Que es la Contraloria General de la Republica?",
             "desc": "Historia, mision y funciones constitucionales de la CGR.",
             "duracion": 15},
            {"num": 2, "titulo": "Estructura organizacional",
             "desc": "Delegadas, Direcciones, Gerencias y el organigrama de la CGR.",
             "duracion": 20},
            {"num": 3, "titulo": "Que es el control fiscal?",
             "desc": "Articulo 267 de la Constitucion y los principios del control fiscal.",
             "duracion": 15},
            {"num": 4, "titulo": "Los macroprocesos: Micro, Macro, Responsabilidad Fiscal, GIA",
             "desc": "Los 4 grandes procesos de la CGR y como se interrelacionan.",
             "duracion": 25},
            {"num": 5, "titulo": "Sistemas institucionales: APA, SIRECI, SIGECI, SIGEDOC",
             "desc": "Las herramientas tecnologicas que usaras en tu dia a dia.",
             "duracion": 20},
            {"num": 6, "titulo": "Tu primer dia como funcionario — que hacer paso a paso",
             "desc": "Guia practica para tus primeras semanas en la CGR.",
             "duracion": 15},
        ],
    },
    {
        "id": "ruta-002",
        "nombre": "Auditoria Financiera DVF",
        "descripcion": "Domina las 5 fases del proceso auditor de la Direccion de Vigilancia Fiscal, con los 30 formatos de la GAF.",
        "icono": "🔍",
        "color": "#1E8449",
        "direccion": "DVF",
        "orden": 2,
        "lecciones": [
            {"num": 1, "titulo": "Que es una auditoria financiera?",
             "desc": "Concepto, objetivos y marco normativo de la auditoria financiera.",
             "duracion": 20},
            {"num": 2, "titulo": "Las 5 fases del proceso auditor",
             "desc": "Pre-planeacion, Planeacion, Ejecucion, Informe y Seguimiento.",
             "duracion": 25},
            {"num": 3, "titulo": "Pre-planeacion: Formatos 1-10",
             "desc": "Conocimiento del sujeto de control, independencia, datos generales.",
             "duracion": 30},
            {"num": 4, "titulo": "Planeacion: Riesgo, materialidad, muestreo",
             "desc": "Formatos 11-20: evaluacion de riesgos y diseno de la auditoria.",
             "duracion": 30},
            {"num": 5, "titulo": "El Formato 14 al detalle: Matriz de Riesgos",
             "desc": "Como construir una matriz de riesgos paso a paso.",
             "duracion": 30},
            {"num": 6, "titulo": "El Formato 17 al detalle: Calculo de Materialidad",
             "desc": "NIA 320 y 450: materialidad global, de ejecucion y umbral.",
             "duracion": 25},
            {"num": 7, "titulo": "Ejecucion: Pruebas de detalle, Benford, hallazgos",
             "desc": "Formatos 21-30: recopilacion de evidencia y pruebas sustantivas.",
             "duracion": 30},
            {"num": 8, "titulo": "El hallazgo: Condicion, Criterio, Causa, Efecto",
             "desc": "Los 4 elementos obligatorios de un hallazgo de auditoria.",
             "duracion": 25},
            {"num": 9, "titulo": "Connotaciones: Administrativa, Fiscal, Disciplinaria, Penal",
             "desc": "Tipos de connotacion, implicaciones y oficios de traslado.",
             "duracion": 20},
            {"num": 10, "titulo": "Informe final, dictamen y seguimiento",
             "desc": "Como se estructura el informe, el dictamen y el seguimiento.",
             "duracion": 25},
        ],
    },
    {
        "id": "ruta-003",
        "nombre": "Estudios Sectoriales DES",
        "descripcion": "Aprende el proceso de control fiscal macro de la Direccion de Estudios Sectoriales.",
        "icono": "📊",
        "color": "#1A5276",
        "direccion": "DES",
        "orden": 3,
        "lecciones": [
            {"num": 1, "titulo": "Que es el control fiscal macro?",
             "desc": "Diferencias con el control micro y el enfoque sectorial.",
             "duracion": 20},
            {"num": 2, "titulo": "Tipos de productos: estudios, EPP, diagnosticos, boletines",
             "desc": "Los diferentes productos que genera la DES.",
             "duracion": 20},
            {"num": 3, "titulo": "Planeacion de un estudio sectorial",
             "desc": "Como se planea un estudio: objetivos, alcance, metodologia.",
             "duracion": 25},
            {"num": 4, "titulo": "Ejecucion: recopilacion y sistematizacion de evidencia",
             "desc": "Fuentes de informacion, bases de datos y sistematizacion.",
             "duracion": 25},
            {"num": 5, "titulo": "Analisis: presupuestal, regalias, politica publica",
             "desc": "Tecnicas de analisis fiscal macro y herramientas.",
             "duracion": 30},
            {"num": 6, "titulo": "Informe final: redaccion, citacion APA/ICONTEC",
             "desc": "Estructura del informe, normas de citacion y revision.",
             "duracion": 20},
            {"num": 7, "titulo": "Seguimiento continuo y Observatorio TIC",
             "desc": "El Observatorio como herramienta de seguimiento sectorial.",
             "duracion": 15},
        ],
    },
    {
        "id": "ruta-004",
        "nombre": "Normativa Esencial",
        "descripcion": "Repasa las leyes, decretos y normas internacionales fundamentales para el control fiscal.",
        "icono": "⚖️",
        "color": "#8E44AD",
        "direccion": "TODOS",
        "orden": 4,
        "lecciones": [
            {"num": 1, "titulo": "Constitucion — Arts 267, 268, 271-274",
             "desc": "El fundamento constitucional del control fiscal en Colombia.",
             "duracion": 20},
            {"num": 2, "titulo": "Ley 42 de 1993 — Organizacion del control fiscal",
             "desc": "Estructura y funcionamiento del control fiscal.",
             "duracion": 20},
            {"num": 3, "titulo": "Ley 610 de 2000 — Proceso de responsabilidad fiscal",
             "desc": "Como funciona el proceso de responsabilidad fiscal.",
             "duracion": 25},
            {"num": 4, "titulo": "Ley 80 de 1993 — Contratacion estatal",
             "desc": "Principios y reglas de la contratacion publica.",
             "duracion": 25},
            {"num": 5, "titulo": "Decreto 403 de 2020 — Principios del control fiscal",
             "desc": "El marco normativo mas reciente del control fiscal.",
             "duracion": 20},
            {"num": 6, "titulo": "ISSAI — Normas internacionales de auditoria",
             "desc": "El marco ISSAI de INTOSAI para las EFS.",
             "duracion": 25},
            {"num": 7, "titulo": "COSO — Marco de control interno",
             "desc": "Los 5 componentes del marco COSO y su aplicacion.",
             "duracion": 25},
            {"num": 8, "titulo": "Ley 1581 de 2012 — Proteccion de datos personales",
             "desc": "Habeas data y tratamiento de informacion personal.",
             "duracion": 15},
        ],
    },
]


# ── Contenido Markdown de ejemplo para lecciones ────────────────────────────

CONTENIDO_LECCION_EJEMPLO: dict[str, str] = {
    "ruta-001-1": """# Que es la Contraloria General de la Republica?

La **Contraloria General de la Republica (CGR)** es el maximo organo de control fiscal del Estado colombiano.

## Fundamento constitucional

Segun el **Articulo 267** de la Constitucion Politica de Colombia:

> *"El control fiscal es una funcion publica que ejercera la Contraloria General de la Republica, la cual vigila la gestion fiscal de la administracion y de los particulares o entidades que manejen fondos o bienes de la Nacion."*

## Funciones principales

1. **Vigilar** la gestion fiscal de entidades publicas
2. **Evaluar** los resultados obtenidos por las entidades
3. **Establecer** la responsabilidad fiscal
4. **Informar** al Congreso sobre el estado de los recursos publicos

## Datos clave

| Aspecto | Detalle |
|---------|---------|
| Creacion | Constitucion de 1991 (Art. 267-274) |
| Sede | Bogota D.C., Colombia |
| Contralor | Elegido por el Congreso |
| Periodo | 4 anos, no reelegible |
| Alcance | Todo el territorio nacional |

## Para que sirve?

Imagina que el Estado es una gran empresa. La CGR es como el **auditor externo** que verifica que el dinero de todos los colombianos se use correctamente.

> 💡 **Ejemplo**: Si el Ministerio de Educacion recibe $5 billones para construir escuelas, la CGR verifica que ese dinero efectivamente se use para eso y no se pierda o desvie.
""",
    "ruta-002-8": """# El Hallazgo: Condicion, Criterio, Causa, Efecto

Un **hallazgo de auditoria** es una diferencia significativa entre lo que *deberia ser* y lo que *realmente es*.

## Los 4 elementos obligatorios

Todo hallazgo DEBE tener estos 4 elementos:

### 1. 🔍 Condicion (Que encontraste?)
Es la **situacion factica** — lo que realmente esta pasando.

> **Ejemplo ficticio**: "La Entidad Ejemplo S.A. suscribio 12 contratos de conectividad por $2.345 millones con valores superiores al promedio del mercado en un 35%."

### 2. ⚖️ Criterio (Contra que norma?)
Es la **norma o regla que se incumplio**. Debe citarse EXACTAMENTE.

> **Ejemplo**: "Ley 80 de 1993, articulo 25, numeral 12 — Principio de economia."

### 3. 🤔 Causa (Por que paso?)
Es la **razon** por la que existe la diferencia.

> **Ejemplo**: "Deficiencia en los estudios de mercado realizados por la entidad."

### 4. 💥 Efecto (Cual es la consecuencia?)
Es el **impacto** — cuantificado cuando sea posible.

> **Ejemplo**: "Presunto detrimento patrimonial estimado en $820 millones."

## Tabla resumen

| Elemento | Pregunta clave | Ejemplo |
|----------|---------------|---------|
| Condicion | Que encontraste? | Sobrecostos del 35% |
| Criterio | Que norma se incumplio? | Ley 80/93, Art. 25 |
| Causa | Por que ocurrio? | Malos estudios de mercado |
| Efecto | Cual es el impacto? | Detrimento de $820M |

## Mini-quiz

A ver, ya sabes los 4 elementos. Piensa:

*Si un municipio contrato una obra por $500 millones pero la obra esta inconclusa al 60%...*

- Cual seria la **condicion**?
- Cual seria el **efecto**?

Piensalo y me cuentas! 😊
""",
}


def _orm_a_dict(obj: Any) -> dict[str, Any]:
    """Convierte un objeto ORM de SQLAlchemy a diccionario serializable."""
    if obj is None:
        return {}
    d: dict[str, Any] = {}
    for col in obj.__table__.columns:
        val = getattr(obj, col.name, None)
        if isinstance(val, datetime):
            val = val.isoformat()
        d[col.name] = val
    return d


class CapacitacionService:
    """Servicio de gestion de capacitacion y rutas de aprendizaje."""

    def __init__(self, db_session: AsyncSession) -> None:
        self._db = db_session

    # ── Rutas ──────────────────────────────────────────────────────────────

    async def listar_rutas(
        self, direccion: Optional[str] = None,
    ) -> list[dict[str, Any]]:
        """Lista rutas de aprendizaje activas, filtradas por direccion."""
        query = select(RutaAprendizaje).where(RutaAprendizaje.activa == True)
        if direccion:
            query = query.where(
                (RutaAprendizaje.direccion == direccion) |
                (RutaAprendizaje.direccion == "TODOS")
            )
        query = query.order_by(RutaAprendizaje.orden)
        resultado = await self._db.execute(query)
        return [_orm_a_dict(r) for r in resultado.scalars().all()]

    async def obtener_ruta(self, ruta_id: str) -> Optional[dict[str, Any]]:
        """Obtiene una ruta por ID."""
        resultado = await self._db.execute(
            select(RutaAprendizaje).where(RutaAprendizaje.id == ruta_id)
        )
        obj = resultado.scalar_one_or_none()
        return _orm_a_dict(obj) if obj else None

    # ── Lecciones ──────────────────────────────────────────────────────────

    async def listar_lecciones(self, ruta_id: str) -> list[dict[str, Any]]:
        """Lista lecciones de una ruta ordenadas."""
        resultado = await self._db.execute(
            select(Leccion)
            .where(Leccion.ruta_id == ruta_id)
            .order_by(Leccion.orden, Leccion.numero)
        )
        return [_orm_a_dict(l) for l in resultado.scalars().all()]

    async def obtener_leccion(self, leccion_id: str) -> Optional[dict[str, Any]]:
        """Obtiene una leccion por ID con contenido Markdown."""
        resultado = await self._db.execute(
            select(Leccion).where(Leccion.id == leccion_id)
        )
        obj = resultado.scalar_one_or_none()
        return _orm_a_dict(obj) if obj else None

    # ── Progreso ──────────────────────────────────────────────────────────

    async def obtener_progreso_usuario(
        self, usuario_id: int, ruta_id: Optional[str] = None,
    ) -> list[dict[str, Any]]:
        """Obtiene el progreso de un usuario, opcionalmente filtrado por ruta."""
        query = select(ProgresoUsuario).where(ProgresoUsuario.usuario_id == usuario_id)
        if ruta_id:
            query = query.where(ProgresoUsuario.ruta_id == ruta_id)
        resultado = await self._db.execute(query)
        return [_orm_a_dict(p) for p in resultado.scalars().all()]

    async def marcar_leccion_completada(
        self, usuario_id: int, ruta_id: str, leccion_id: str,
    ) -> dict[str, Any]:
        """Marca una leccion como completada."""
        # Buscar progreso existente
        resultado = await self._db.execute(
            select(ProgresoUsuario).where(
                ProgresoUsuario.usuario_id == usuario_id,
                ProgresoUsuario.leccion_id == leccion_id,
            )
        )
        progreso = resultado.scalar_one_or_none()

        if progreso:
            progreso.completada = True
            progreso.fecha_completada = datetime.now(timezone.utc)
        else:
            progreso = ProgresoUsuario(
                id=str(uuid.uuid4()),
                usuario_id=usuario_id,
                ruta_id=ruta_id,
                leccion_id=leccion_id,
                completada=True,
                fecha_completada=datetime.now(timezone.utc),
            )
            self._db.add(progreso)

        await self._db.flush()
        return _orm_a_dict(progreso)

    async def obtener_resumen_progreso(
        self, usuario_id: int,
    ) -> dict[str, Any]:
        """Resumen de progreso global del usuario."""
        rutas = await self.listar_rutas()
        progreso_total = await self.obtener_progreso_usuario(usuario_id)

        completadas_por_ruta: dict[str, int] = {}
        for p in progreso_total:
            if p.get("completada"):
                rid = p.get("ruta_id", "")
                completadas_por_ruta[rid] = completadas_por_ruta.get(rid, 0) + 1

        resumen_rutas = []
        total_lecciones = 0
        total_completadas = 0

        for ruta in rutas:
            ruta_id = ruta["id"]
            lecciones = await self.listar_lecciones(ruta_id)
            n_lecciones = len(lecciones)
            n_completadas = completadas_por_ruta.get(ruta_id, 0)
            total_lecciones += n_lecciones
            total_completadas += n_completadas

            resumen_rutas.append({
                "ruta_id": ruta_id,
                "nombre": ruta["nombre"],
                "icono": ruta.get("icono", ""),
                "color": ruta.get("color", ""),
                "direccion": ruta.get("direccion", ""),
                "total_lecciones": n_lecciones,
                "completadas": n_completadas,
                "porcentaje": round(n_completadas / n_lecciones * 100) if n_lecciones > 0 else 0,
            })

        # Quizzes
        resultado_quizzes = await self._db.execute(
            select(QuizResultado).where(QuizResultado.usuario_id == usuario_id)
        )
        quizzes = list(resultado_quizzes.scalars().all())
        quizzes_aprobados = sum(1 for q in quizzes if q.aprobado)

        return {
            "total_rutas": len(rutas),
            "total_lecciones": total_lecciones,
            "lecciones_completadas": total_completadas,
            "porcentaje_global": round(total_completadas / total_lecciones * 100) if total_lecciones > 0 else 0,
            "quizzes_realizados": len(quizzes),
            "quizzes_aprobados": quizzes_aprobados,
            "rutas": resumen_rutas,
        }

    # ── Quizzes ─────────────────────────────────────────────────────────────

    async def registrar_quiz(
        self,
        usuario_id: int,
        ruta_id: str,
        leccion_id: Optional[str],
        puntaje: float,
        total_preguntas: int,
        respuestas: list[dict[str, Any]],
    ) -> dict[str, Any]:
        """Registra el resultado de un quiz."""
        aprobado = puntaje >= 70.0

        quiz = QuizResultado(
            id=str(uuid.uuid4()),
            usuario_id=usuario_id,
            ruta_id=ruta_id,
            leccion_id=leccion_id,
            puntaje=puntaje,
            total_preguntas=total_preguntas,
            respuestas_json=respuestas,
            aprobado=aprobado,
        )
        self._db.add(quiz)
        await self._db.flush()

        logger.info(
            "Quiz ruta %s usuario %d: %.1f%% (%s)",
            ruta_id, usuario_id, puntaje, "APROBADO" if aprobado else "NO APROBADO",
        )
        return _orm_a_dict(quiz)

    async def obtener_quizzes_usuario(
        self, usuario_id: int, ruta_id: Optional[str] = None,
    ) -> list[dict[str, Any]]:
        """Lista quizzes de un usuario."""
        query = select(QuizResultado).where(QuizResultado.usuario_id == usuario_id)
        if ruta_id:
            query = query.where(QuizResultado.ruta_id == ruta_id)
        query = query.order_by(QuizResultado.created_at.desc())
        resultado = await self._db.execute(query)
        return [_orm_a_dict(q) for q in resultado.scalars().all()]

    # ── Metricas (solo directores/admin) ────────────────────────────────────

    async def obtener_metricas_globales(self) -> dict[str, Any]:
        """Metricas globales de capacitacion."""
        # Usuarios con progreso
        resultado_usuarios = await self._db.execute(
            select(func.count(func.distinct(ProgresoUsuario.usuario_id)))
        )
        funcionarios_activos = resultado_usuarios.scalar() or 0

        # Lecciones completadas
        resultado_completadas = await self._db.execute(
            select(func.count(ProgresoUsuario.id))
            .where(ProgresoUsuario.completada == True)
        )
        total_completadas = resultado_completadas.scalar() or 0

        # Quizzes
        resultado_quizzes = await self._db.execute(
            select(
                func.count(QuizResultado.id),
                func.avg(QuizResultado.puntaje),
                func.count(QuizResultado.id).filter(QuizResultado.aprobado == True),
            )
        )
        row = resultado_quizzes.one()
        total_quizzes = row[0] or 0
        puntaje_promedio = float(row[1]) if row[1] else 0
        quizzes_aprobados = row[2] or 0

        # Progreso por ruta
        rutas = await self.listar_rutas()
        progreso_rutas = []
        for ruta in rutas:
            res = await self._db.execute(
                select(func.count(func.distinct(ProgresoUsuario.usuario_id)))
                .where(
                    ProgresoUsuario.ruta_id == ruta.id,
                    ProgresoUsuario.completada == True,
                )
            )
            n_usuarios = res.scalar() or 0
            progreso_rutas.append({
                "ruta": ruta.nombre,
                "direccion": ruta.direccion,
                "usuarios_completaron": n_usuarios,
            })

        return {
            "funcionarios_activos": funcionarios_activos,
            "total_lecciones_completadas": total_completadas,
            "total_quizzes": total_quizzes,
            "puntaje_promedio": round(puntaje_promedio, 1),
            "quizzes_aprobados": quizzes_aprobados,
            "tasa_aprobacion": round(quizzes_aprobados / total_quizzes * 100, 1) if total_quizzes > 0 else 0,
            "progreso_por_ruta": progreso_rutas,
        }

    # ── Seed ────────────────────────────────────────────────────────────────

    async def seed_rutas_y_lecciones(self) -> int:
        """Inserta las rutas y lecciones de semilla. Retorna total creados."""
        creados = 0
        for ruta_data in RUTAS_SEED:
            # Verificar si ya existe
            existente = await self.obtener_ruta(ruta_data["id"])
            if existente:
                continue

            ruta = RutaAprendizaje(
                id=ruta_data["id"],
                nombre=ruta_data["nombre"],
                descripcion=ruta_data["descripcion"],
                icono=ruta_data["icono"],
                color=ruta_data["color"],
                direccion=ruta_data["direccion"],
                orden=ruta_data["orden"],
                activa=True,
                total_lecciones=len(ruta_data["lecciones"]),
            )
            self._db.add(ruta)

            for leccion_data in ruta_data["lecciones"]:
                leccion_id = f"{ruta_data['id']}-{leccion_data['num']}"
                contenido = CONTENIDO_LECCION_EJEMPLO.get(
                    f"{ruta_data['id']}-{leccion_data['num']}",
                    f"# {leccion_data['titulo']}\n\n"
                    f"{leccion_data['desc']}\n\n"
                    f"*Contenido detallado disponible proximamente. "
                    f"Mientras tanto, preguntale a CecilIA en el chat!*",
                )

                leccion = Leccion(
                    id=leccion_id,
                    ruta_id=ruta_data["id"],
                    numero=leccion_data["num"],
                    titulo=leccion_data["titulo"],
                    descripcion=leccion_data["desc"],
                    contenido_md=contenido,
                    duracion_minutos=leccion_data["duracion"],
                    orden=leccion_data["num"],
                )
                self._db.add(leccion)
                creados += 1

            creados += 1  # la ruta

        await self._db.flush()
        logger.info("Seed de capacitacion: %d registros creados", creados)
        return creados


# ── Generadores de contenido didactico ───────────────────────────────────────


class _ManualDidacticoDocx(FormatoBaseCGR):
    """Genera un manual didactico profesional en DOCX con logos CGR y CecilIA."""

    def __init__(self, tema: str, nivel: str, contenido: str) -> None:
        self._tema = tema
        self._nivel = nivel
        self._contenido = contenido
        super().__init__(
            numero_formato=0,
            nombre_formato=f"Manual Didactico — {tema}",
            datos={
                "nombre_entidad": "Contraloria General de la Republica",
                "vigencia": datetime.now().strftime("%Y"),
            },
        )

    def _agregar_portada(self) -> None:
        """Agrega portada profesional con logos CGR y CecilIA."""
        from pathlib import Path
        from docx.shared import Cm, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Logo CGR
        self._agregar_logo()

        # Logo CecilIA
        rutas_cecilia = [
            Path(__file__).parent.parent / "static" / "logo-cecilia.png",
            Path("/app/app/static/logo-cecilia.png"),
        ]
        for ruta in rutas_cecilia:
            if ruta.exists():
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(12)
                run = p.add_run()
                run.add_picture(str(ruta), width=Cm(3))
                break

        # Titulo del manual
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(20)
        run = p.add_run(f"MANUAL DIDACTICO")
        run.font.name = "Georgia"
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)

        p2 = self.doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(self._tema.upper())
        run2.font.name = "Georgia"
        run2.font.size = Pt(16)
        run2.font.bold = True
        run2.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)

        # Subtitulo
        p3 = self.doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_before = Pt(10)
        run3 = p3.add_run(f"Nivel: {self._nivel.title()} — Material de Capacitacion")
        run3.font.name = "Calibri"
        run3.font.size = Pt(12)
        run3.font.color.rgb = RGBColor(0x5F, 0x63, 0x68)

        # Linea dorada
        self._agregar_linea_dorada()

        # Creditos
        p4 = self.doc.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p4.paragraph_format.space_before = Pt(30)
        run4 = p4.add_run("Contraloria General de la Republica")
        run4.font.name = "Calibri"
        run4.font.size = Pt(11)
        run4.font.bold = True

        p5 = self.doc.add_paragraph()
        p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run5 = p5.add_run("Contraloria Delegada para el Sector TIC — CD-TIC")
        run5.font.name = "Calibri"
        run5.font.size = Pt(10)

        p6 = self.doc.add_paragraph()
        p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p6.paragraph_format.space_before = Pt(8)
        run6 = p6.add_run("Proyecto concebido e impulsado por el Dr. Omar Javier Contreras Socarras")
        run6.font.name = "Calibri"
        run6.font.size = Pt(9)
        run6.font.italic = True
        run6.font.color.rgb = RGBColor(0x5F, 0x63, 0x68)

        p7 = self.doc.add_paragraph()
        p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p7.paragraph_format.space_before = Pt(15)
        run7 = p7.add_run(f"Version 1.0 — {datetime.now().strftime('%B %Y').title()}")
        run7.font.name = "Calibri"
        run7.font.size = Pt(10)
        run7.font.color.rgb = RGBColor(0xC9, 0xA8, 0x4C)

        # Salto de pagina
        self.doc.add_page_break()

    def construir(self) -> None:
        # Portada profesional
        self._agregar_portada()

        # Encabezado institucional en la segunda pagina
        self.agregar_encabezado()

        # Tabla de metadata
        self.crear_tabla_clave_valor([
            ("Tema", self._tema),
            ("Nivel", self._nivel.title()),
            ("Tipo", "Material de capacitacion — CecilIA v2"),
            ("Fecha", datetime.now().strftime("%d/%m/%Y")),
        ])

        # Renderizar contenido Markdown como secciones DOCX
        lineas = self._contenido.split("\n")
        for linea in lineas:
            linea_strip = linea.strip()
            if not linea_strip:
                continue
            if linea_strip.startswith("## "):
                self.agregar_titulo_seccion(linea_strip[3:])
            elif linea_strip.startswith("# "):
                self.agregar_titulo_seccion(linea_strip[2:])
            elif linea_strip.startswith("- "):
                self.agregar_parrafo_justificado(f"  • {linea_strip[2:]}")
            elif linea_strip.startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.")):
                self.agregar_parrafo_justificado(f"  {linea_strip}")
            elif linea_strip.startswith(">"):
                self.agregar_parrafo_justificado(linea_strip[1:].strip())
            else:
                # Limpiar marcadores **bold**
                texto_limpio = linea_strip.replace("**", "")
                self.agregar_parrafo_justificado(texto_limpio)

        self.agregar_nota_ia()


def generar_manual_docx(tema: str, nivel: str = "basico", llm=None) -> bytes:
    """Genera un manual didactico profesional en formato DOCX.

    Si el LLM esta disponible, genera contenido especializado.
    Si no, usa contenido de ejemplo mejorado.
    """
    contenido = None

    # Intentar generar contenido con LLM
    if llm:
        try:
            import asyncio
            prompt = (
                f"Eres un experto en control fiscal colombiano y pedagogia. "
                f"Genera un manual didactico de nivel {nivel} sobre: '{tema}'. "
                f"El manual es para capacitacion de funcionarios de la Contraloria General de la Republica (CGR). "
                f"\nEstructura requerida (usa formato Markdown con ## para secciones):\n"
                f"## Introduccion\n(Contexto del tema en el control fiscal colombiano)\n"
                f"## Marco normativo\n(Leyes, decretos y normas aplicables: cita Ley 42/1993, Ley 610/2000, Decreto 403/2020, ISSAI, etc.)\n"
                f"## Conceptos clave\n(Definiciones y explicaciones detalladas)\n"
                f"## Procedimiento paso a paso\n(Como se aplica en la practica auditora)\n"
                f"## Ejemplo practico\n(Caso con datos FICTICIOS de una entidad ejemplo)\n"
                f"## Errores frecuentes\n(Los 5-7 errores mas comunes y como evitarlos)\n"
                f"## Recomendaciones del auditor experimentado\n(Tips practicos)\n"
                f"## Normativa de referencia\n(Lista completa de normas aplicables)\n"
                f"\nIMPORTANTE: Todos los datos y ejemplos deben ser FICTICIOS. "
                f"Cita normatividad REAL colombiana. Escribe en espanol."
            )
            resp = asyncio.get_event_loop().run_until_complete(
                asyncio.to_thread(llm.invoke, prompt)
            ) if asyncio.get_event_loop().is_running() else llm.invoke(prompt)
            contenido = resp.content if hasattr(resp, "content") else str(resp)
        except Exception as e:
            logger.warning("No se pudo generar contenido con LLM para manual: %s", e)

    if not contenido:
        contenido = f"""## Introduccion

Este manual explica de forma clara y sencilla el tema de {tema} para funcionarios de la Contraloria General de la Republica en proceso de capacitacion.

Todos los datos y ejemplos son ficticios y se usan unicamente con fines educativos, conforme a la Circular 023 de la CGR sobre uso responsable de inteligencia artificial.

## Marco normativo

El control fiscal en Colombia se fundamenta en las siguientes normas:

- Constitucion Politica de Colombia, Articulos 267 a 274: Establece las funciones de la Contraloria General de la Republica
- Ley 42 de 1993: Organiza el sistema de control fiscal financiero
- Ley 610 de 2000: Regula el proceso de responsabilidad fiscal
- Decreto 403 de 2020: Desarrolla los principios del control fiscal
- ISSAI 100, 200, 300 y 400: Normas internacionales de auditoria del sector publico
- Circular 023 de la CGR: Orientaciones sobre el uso de inteligencia artificial

## Conceptos clave

{tema} es un concepto fundamental en el control fiscal colombiano. Comprende los procesos, procedimientos y metodologias que los auditores de la CGR aplican para vigilar la gestion de los recursos publicos.

Los auditores deben entender este tema para cumplir con su mision constitucional de proteger el patrimonio publico de todos los colombianos.

## Procedimiento paso a paso

1. Identificar el alcance del tema dentro de la auditoria
2. Revisar la normatividad aplicable
3. Recopilar evidencia documental
4. Analizar la informacion recopilada
5. Documentar los hallazgos siguiendo los 4 elementos (condicion, criterio, causa, efecto)
6. Comunicar resultados a la entidad auditada
7. Hacer seguimiento a los planes de mejoramiento

## Ejemplo practico

Supongamos que estamos auditando a la Entidad Ejemplo S.A. (NIT 900.000.001-1) con un presupuesto de $50.000 millones para la vigencia 2025.

Al revisar la ejecucion presupuestal, encontramos que el 15% de los recursos asignados al rubro de inversion no tienen soporte documental adecuado. Esto constituye un posible hallazgo con connotacion administrativa y fiscal.

## Errores frecuentes

1. No documentar adecuadamente las fuentes de informacion
2. Usar datos desactualizados o sin verificar
3. No verificar la coherencia de las cifras reportadas
4. Omitir la fundamentacion normativa en los hallazgos
5. No comunicar oportunamente los resultados preliminares
6. Confundir observaciones con hallazgos
7. No aplicar el principio de materialidad

## Recomendaciones del auditor experimentado

- Siempre verifica contra la norma original, no dependas de interpretaciones de terceros
- Documenta todo, incluso lo que parece obvio en el momento
- Pide ayuda cuando no entiendas algo, la colaboracion entre auditores es fundamental
- Usa CecilIA para resolver dudas rapidas sobre normatividad y procedimientos
- Mantente actualizado con las circulares y directivas de la CGR

## Normativa de referencia

- Constitucion Politica de Colombia, Art. 267-274
- Ley 42 de 1993 (Control fiscal financiero)
- Ley 610 de 2000 (Responsabilidad fiscal)
- Decreto 403 de 2020 (Principios del control fiscal)
- Decreto 2037 de 2019 (Estructura de la CGR)
- ISSAI 100, 200, 300, 400 (Normas internacionales)
- Guia de Auditoria Financiera de la CGR (GAF)
- Circular 023 de la CGR (Uso de IA)
"""

    generador = _ManualDidacticoDocx(tema, nivel, contenido)
    return generador.generar_bytes()


def generar_podcast_script(tema: str) -> str:
    """Genera un script conversacional tipo podcast."""
    return f"""# Script de Podcast: {tema}
## CecilIA Explica — Episodio de Capacitacion

**[INTRO - Musica suave institucional]**

**Narrador**: Bienvenidos a "CecilIA Explica", el podcast de capacitacion de la Contraloria General de la Republica. Hoy hablaremos sobre: **{tema}**.

**Narrador**: Nos acompana Ana, una auditora que acaba de ingresar a la CGR. Ana, bienvenida!

**Ana**: Gracias! Estoy emocionada de aprender. Tengo muchas preguntas sobre {tema.lower()}.

**Narrador**: Perfecto, empecemos desde lo basico.

---

**Narrador**: Ana, sabes que es {tema.lower()}?

**Ana**: Tengo una idea general, pero no estoy segura de los detalles tecnicos...

**Narrador**: No te preocupes, vamos paso a paso. {tema} es un concepto fundamental en el control fiscal. Imagina que eres una detective financiera...

**Ana**: Ah, eso suena interesante! Como asi?

**Narrador**: Bueno, tu trabajo es verificar que el dinero publico se use correctamente. Y para eso necesitas entender {tema.lower()}, que es basicamente tu "lupa" para encontrar irregularidades.

**Ana**: Entonces, en la practica, como funciona?

**Narrador**: Veamos un ejemplo con datos ficticios. Supongamos que la Entidad Ejemplo S.A. recibio $50.000 millones para un proyecto de infraestructura...

**Ana**: Y si encuentro algo raro?

**Narrador**: Excelente pregunta! Ahi es donde entra el concepto del hallazgo. Tienes que documentar 4 elementos: Condicion, Criterio, Causa y Efecto.

**Ana**: Cuatro elementos... eso es mucho!

**Narrador**: Parece mucho, pero con practica se vuelve natural. CecilIA puede ayudarte a estructurarlos.

---

**[CIERRE]**

**Narrador**: Gracias por escucharnos. Recuerda: todos los datos fueron ficticios, con fines educativos.

**Ana**: Y si tienen dudas, preguntale a CecilIA!

**Narrador**: Hasta la proxima!

**[Musica de cierre]**

---
*📚 Material de capacitacion — Datos ficticios de ejemplo. CecilIA — Circular 023 CGR*
"""


def generar_glosario_docx() -> bytes:
    """Genera un glosario de terminos de control fiscal en DOCX."""
    class _Glosario(FormatoBaseCGR):
        def __init__(self):
            super().__init__(
                numero_formato=0,
                nombre_formato="Glosario de Control Fiscal",
                datos={
                    "nombre_entidad": "Contraloria General de la Republica",
                    "vigencia": datetime.now().strftime("%Y"),
                },
            )

        def construir(self):
            self.agregar_titulo_seccion("Glosario de Terminos de Control Fiscal")
            self.agregar_parrafo_justificado(
                "Este glosario contiene los terminos mas utilizados en el control fiscal colombiano. "
                "Material de referencia para funcionarios en proceso de capacitacion."
            )

            terminos = [
                ("Auditoria", "Proceso sistematico de evaluacion de la gestion fiscal de una entidad."),
                ("Benford (Ley de)", "Distribucion estadistica que predice la frecuencia del primer digito en conjuntos de datos numericos."),
                ("CGR", "Contraloria General de la Republica de Colombia."),
                ("Connotacion", "Tipo de implicacion de un hallazgo: administrativa, fiscal, disciplinaria o penal."),
                ("Control fiscal", "Funcion publica de vigilancia de la gestion fiscal del Estado (Art. 267 CN)."),
                ("COSO", "Committee of Sponsoring Organizations — Marco de control interno con 5 componentes."),
                ("Criterio", "Norma, ley o regla contra la cual se evalua la condicion encontrada."),
                ("Cuantia", "Valor monetario estimado del presunto dano patrimonial."),
                ("DES", "Direccion de Estudios Sectoriales — control fiscal macro."),
                ("Dictamen", "Opinion del auditor sobre los estados financieros de la entidad."),
                ("DVF", "Direccion de Vigilancia Fiscal — control fiscal micro."),
                ("Efecto", "Consecuencia o impacto de la irregularidad encontrada."),
                ("Evidencia", "Documentos, registros o pruebas que soportan un hallazgo."),
                ("GAF", "Guia de Auditoria Financiera de la CGR — contiene los 30 formatos."),
                ("GIA", "Grupo de Investigaciones Administrativas."),
                ("Hallazgo", "Diferencia significativa entre lo esperado y lo encontrado en una auditoria."),
                ("INTOSAI", "International Organization of Supreme Audit Institutions."),
                ("ISSAI", "International Standards of Supreme Audit Institutions."),
                ("Materialidad", "Umbral de importancia relativa para determinar errores significativos (NIA 320)."),
                ("Muestreo", "Seleccion de una muestra representativa para pruebas de auditoria."),
                ("NIA", "Normas Internacionales de Auditoria."),
                ("Observacion", "Comunicacion de hallazgos preliminares a la entidad auditada."),
                ("Plan de mejoramiento", "Acciones correctivas que la entidad se compromete a implementar."),
                ("Presunto dano patrimonial", "Perdida o perjuicio economico al Estado pendiente de determinacion."),
                ("Responsabilidad fiscal", "Proceso para determinar dano patrimonial y obtener resarcimiento."),
                ("Riesgo de auditoria", "Probabilidad de emitir una opinion incorrecta sobre los estados financieros."),
                ("Riesgo inherente", "Riesgo propio de la naturaleza de la transaccion o cuenta."),
                ("Riesgo de control", "Riesgo de que el control interno no prevenga o detecte errores."),
                ("SECOP", "Sistema Electronico de Contratacion Publica."),
                ("SIGECI", "Sistema Integrado de Gestion de la CGR."),
                ("SIRECI", "Sistema de Rendicion Electronica de Cuenta e Informes."),
                ("Sujeto de control", "Entidad publica objeto de la auditoria."),
                ("Traslado", "Envio de hallazgos con connotacion a la entidad competente."),
            ]

            encabezados = ["Termino", "Definicion"]
            filas = [[t, d] for t, d in terminos]
            self.crear_tabla(encabezados=encabezados, filas=filas)

            self.agregar_nota_ia()

    gen = _Glosario()
    return gen.generar_bytes()


def generar_guia_formato_docx(numero_formato: int) -> bytes:
    """Genera guia detallada de un formato CGR especifico."""
    from app.formatos.registro import CATALOGO_FORMATOS

    info = CATALOGO_FORMATOS.get(numero_formato, {})
    nombre = info.get("nombre", f"Formato F{numero_formato:02d}")
    fase = info.get("fase", "no especificada")

    class _GuiaFormato(FormatoBaseCGR):
        def __init__(self):
            super().__init__(
                numero_formato=0,
                nombre_formato=f"Guia Didactica — {nombre}",
                datos={
                    "nombre_entidad": "Contraloria General de la Republica",
                    "vigencia": datetime.now().strftime("%Y"),
                },
            )

        def construir(self):
            self.agregar_titulo_seccion(f"Guia Didactica: Formato F{numero_formato:02d}")
            self.agregar_titulo_seccion(nombre)

            self.crear_tabla_clave_valor([
                ("Numero", f"F{numero_formato:02d}"),
                ("Nombre", nombre),
                ("Fase", fase.title()),
                ("Tipo", "Material de capacitacion"),
            ])

            self.agregar_titulo_seccion("1. Proposito")
            self.agregar_parrafo_justificado(
                f"El Formato F{numero_formato:02d} ({nombre}) se utiliza durante la fase "
                f"de {fase} del proceso auditor. Su proposito es documentar de forma "
                f"estandarizada la informacion requerida en esta etapa."
            )

            self.agregar_titulo_seccion("2. Cuando se usa?")
            self.agregar_parrafo_justificado(
                f"Se diligencia durante la fase de {fase}. El auditor debe completarlo "
                f"antes de avanzar a la siguiente etapa del proceso."
            )

            self.agregar_titulo_seccion("3. Como se llena? (paso a paso)")
            self.agregar_parrafo_justificado(
                "1. Abra el formato desde CecilIA o descargue la plantilla DOCX\n"
                "2. Complete los datos de la entidad auditada (nombre, NIT, vigencia)\n"
                "3. Diligencie cada seccion siguiendo las instrucciones del formato\n"
                "4. Revise la coherencia de la informacion ingresada\n"
                "5. Solicite validacion del supervisor/coordinador"
            )

            self.agregar_titulo_seccion("4. Ejemplo con datos ficticios")
            self.crear_tabla_clave_valor([
                ("Entidad", "Entidad Ejemplo S.A."),
                ("NIT", "900.000.001-1"),
                ("Vigencia", "2025"),
                ("Tipo de auditoria", "Financiera y de Cumplimiento"),
                ("Presupuesto", "$50.000.000.000"),
            ])

            self.agregar_titulo_seccion("5. Errores comunes")
            self.agregar_parrafo_justificado(
                "- No completar todos los campos obligatorios\n"
                "- Usar datos de vigencias anteriores sin actualizar\n"
                "- No fundamentar con la norma correcta\n"
                "- Omitir la firma del auditor y supervisor\n"
                "- No adjuntar los soportes requeridos"
            )

            self.agregar_titulo_seccion("6. Tips")
            self.agregar_parrafo_justificado(
                "- Usa CecilIA para pre-diligenciar el formato con datos del RAG\n"
                "- Siempre verifica contra la fuente original\n"
                "- Guarda copias de respaldo de cada version\n"
                "- Consulta al supervisor ante cualquier duda"
            )

            self.agregar_nota_ia()

    gen = _GuiaFormato()
    return gen.generar_bytes()
