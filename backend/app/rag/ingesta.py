"""
CecilIA v2 — Sistema de IA para Control Fiscal
Contraloría General de la República de Colombia

Archivo: ingesta.py
Propósito: Ingestión de documentos — PDF (PyMuPDF), DOCX, XLSX, OCR (Tesseract)
Sprint: 0
Autor: Equipo Técnico CecilIA — CD-TIC-CGR
Fecha: Abril 2026
"""

from __future__ import annotations

import logging
import tempfile
from pathlib import Path
from typing import Any, Optional

logger = logging.getLogger("cecilia.rag.ingesta")

# Extensiones soportadas
EXTENSIONES_SOPORTADAS: set[str] = {".pdf", ".docx", ".xlsx", ".xls", ".png", ".jpg", ".jpeg", ".tiff", ".bmp"}


class DocumentoIngestado:
    """Representa un documento procesado y listo para chunking."""

    def __init__(
        self,
        contenido: str,
        metadata: dict[str, Any],
        nombre_archivo: str,
        tipo: str,
        paginas: int = 1,
    ) -> None:
        self.contenido: str = contenido
        self.metadata: dict[str, Any] = metadata
        self.nombre_archivo: str = nombre_archivo
        self.tipo: str = tipo
        self.paginas: int = paginas

    def __repr__(self) -> str:
        return (
            f"DocumentoIngestado(nombre='{self.nombre_archivo}', "
            f"tipo='{self.tipo}', paginas={self.paginas}, "
            f"caracteres={len(self.contenido)})"
        )


def _extraer_texto_pdf(ruta: Path) -> tuple[str, int]:
    """Extrae texto de un archivo PDF usando PyMuPDF (fitz).

    Args:
        ruta: Ruta al archivo PDF.

    Returns:
        Tupla (texto_completo, numero_de_paginas).
    """
    import fitz  # PyMuPDF

    texto_paginas: list[str] = []
    num_paginas: int = 0

    with fitz.open(str(ruta)) as documento:
        num_paginas = len(documento)
        for pagina in documento:
            texto: str = pagina.get_text("text")
            if texto.strip():
                texto_paginas.append(texto)
            else:
                # Página sin texto extraíble — posible imagen/escaneado
                logger.info(
                    "Página %d de '%s' sin texto extraíble; se intentará OCR.",
                    pagina.number + 1,
                    ruta.name,
                )
                texto_ocr: str = _aplicar_ocr_a_pagina(pagina)
                if texto_ocr.strip():
                    texto_paginas.append(texto_ocr)

    return "\n\n".join(texto_paginas), num_paginas


def _aplicar_ocr_a_pagina(pagina: Any) -> str:
    """Aplica OCR a una página de PDF usando Tesseract.

    Args:
        pagina: Objeto página de PyMuPDF.

    Returns:
        Texto reconocido por OCR.
    """
    try:
        import pytesseract
        from PIL import Image
        import io

        # Renderizar página como imagen
        pixmap = pagina.get_pixmap(dpi=300)
        imagen_bytes: bytes = pixmap.tobytes("png")
        imagen = Image.open(io.BytesIO(imagen_bytes))

        texto: str = pytesseract.image_to_string(imagen, lang="spa")
        return texto

    except ImportError:
        logger.warning("pytesseract o PIL no disponibles para OCR.")
        return ""
    except Exception:
        logger.exception("Error al aplicar OCR a página.")
        return ""


def _extraer_texto_docx(ruta: Path) -> str:
    """Extrae texto de un archivo DOCX usando python-docx.

    Args:
        ruta: Ruta al archivo DOCX.

    Returns:
        Texto completo del documento.
    """
    from docx import Document

    documento = Document(str(ruta))
    parrafos: list[str] = []

    for parrafo in documento.paragraphs:
        if parrafo.text.strip():
            parrafos.append(parrafo.text)

    # Extraer texto de tablas
    for tabla in documento.tables:
        for fila in tabla.rows:
            celdas: list[str] = [celda.text.strip() for celda in fila.cells if celda.text.strip()]
            if celdas:
                parrafos.append(" | ".join(celdas))

    return "\n\n".join(parrafos)


def _extraer_texto_xlsx(ruta: Path) -> str:
    """Extrae texto de un archivo XLSX usando openpyxl.

    Args:
        ruta: Ruta al archivo XLSX.

    Returns:
        Texto tabulado del archivo.
    """
    from openpyxl import load_workbook

    wb = load_workbook(str(ruta), data_only=True, read_only=True)
    hojas_texto: list[str] = []

    for nombre_hoja in wb.sheetnames:
        hoja = wb[nombre_hoja]
        filas_texto: list[str] = [f"=== Hoja: {nombre_hoja} ==="]

        for fila in hoja.iter_rows(values_only=True):
            valores: list[str] = [
                str(celda) if celda is not None else ""
                for celda in fila
            ]
            linea: str = " | ".join(valores).strip()
            if linea.replace("|", "").strip():
                filas_texto.append(linea)

        hojas_texto.append("\n".join(filas_texto))

    wb.close()
    return "\n\n".join(hojas_texto)


def _extraer_texto_imagen(ruta: Path) -> str:
    """Aplica OCR a una imagen usando Tesseract.

    Args:
        ruta: Ruta al archivo de imagen.

    Returns:
        Texto reconocido por OCR.
    """
    try:
        import pytesseract
        from PIL import Image

        imagen = Image.open(str(ruta))
        texto: str = pytesseract.image_to_string(imagen, lang="spa")
        return texto

    except ImportError:
        logger.warning("pytesseract o PIL no disponibles para OCR.")
        return ""
    except Exception:
        logger.exception("Error al aplicar OCR a imagen '%s'.", ruta.name)
        return ""


def ingestar_documento(
    ruta: Path | str,
    metadata_extra: Optional[dict[str, Any]] = None,
    coleccion: str = "general",
) -> DocumentoIngestado:
    """Ingesta un documento y extrae su contenido textual.

    Soporta PDF, DOCX, XLSX, e imágenes (OCR).

    Args:
        ruta: Ruta al archivo a ingestar.
        metadata_extra: Metadatos adicionales para asociar al documento.
        coleccion: Nombre de la colección en la base vectorial.

    Returns:
        Documento ingestado con contenido y metadatos.

    Raises:
        ValueError: Si el formato no está soportado.
        FileNotFoundError: Si el archivo no existe.
    """
    ruta = Path(ruta)

    if not ruta.exists():
        raise FileNotFoundError(f"Archivo no encontrado: {ruta}")

    extension: str = ruta.suffix.lower()
    if extension not in EXTENSIONES_SOPORTADAS:
        raise ValueError(
            f"Formato '{extension}' no soportado. "
            f"Formatos válidos: {', '.join(sorted(EXTENSIONES_SOPORTADAS))}"
        )

    logger.info("Ingestando documento: %s (tipo: %s)", ruta.name, extension)

    contenido: str = ""
    paginas: int = 1
    tipo: str = extension.lstrip(".")

    if extension == ".pdf":
        contenido, paginas = _extraer_texto_pdf(ruta)
    elif extension == ".docx":
        contenido = _extraer_texto_docx(ruta)
    elif extension in {".xlsx", ".xls"}:
        contenido = _extraer_texto_xlsx(ruta)
    elif extension in {".png", ".jpg", ".jpeg", ".tiff", ".bmp"}:
        contenido = _extraer_texto_imagen(ruta)
        tipo = "imagen_ocr"

    if not contenido.strip():
        logger.warning("No se pudo extraer texto de '%s'.", ruta.name)

    metadata: dict[str, Any] = {
        "nombre_archivo": ruta.name,
        "ruta_original": str(ruta),
        "tipo": tipo,
        "coleccion": coleccion,
        "paginas": paginas,
        "tamano_bytes": ruta.stat().st_size,
    }
    if metadata_extra:
        metadata.update(metadata_extra)

    documento = DocumentoIngestado(
        contenido=contenido,
        metadata=metadata,
        nombre_archivo=ruta.name,
        tipo=tipo,
        paginas=paginas,
    )

    logger.info("Documento ingestado: %s", documento)
    return documento
