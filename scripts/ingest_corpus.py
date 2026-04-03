"""
CecilIA v2 — Sistema de IA para Control Fiscal
Contraloria General de la Republica de Colombia

Archivo: ingest_corpus.py
Proposito: Ingestion del corpus documental — escaneo de 7 colecciones,
           extraccion de texto, chunking (con modo juridico), generacion
           de embeddings y almacenamiento en pgvector (fragmentos_vectoriales)
Sprint: 1
Autor: Equipo Tecnico CecilIA — CD-TIC-CGR
Fecha: Abril 2026

Uso dentro de Docker:
    docker compose exec backend python scripts/ingest_corpus.py

Uso local (requiere DATABASE_URL_SYNC configurado):
    python scripts/ingest_corpus.py
"""

import hashlib
import os
import sys
import time
import logging
import uuid
from pathlib import Path
from dataclasses import dataclass, field
from typing import Iterator

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)
logger = logging.getLogger("ingest_corpus")

# ---------------------------------------------------------------------------
# Configuracion
# ---------------------------------------------------------------------------

# URL de la base de datos (sincrona para este script)
DATABASE_URL: str = os.environ.get("DATABASE_URL_SYNC", "")
if not DATABASE_URL:
    # Intentar convertir la URL asincrona
    _url = os.environ.get("DATABASE_URL", "")
    if _url:
        DATABASE_URL = _url.replace("postgresql+asyncpg://", "postgresql+psycopg2://")
        if DATABASE_URL.startswith("postgresql://"):
            DATABASE_URL = DATABASE_URL.replace("postgresql://", "postgresql+psycopg2://", 1)

CORPUS_DIR: Path = Path(os.environ.get("CORPUS_DIR", "/datos/corpus"))
CHUNK_SIZE: int = int(os.environ.get("RAG_CHUNK_SIZE", "1000"))
CHUNK_OVERLAP: int = int(os.environ.get("RAG_CHUNK_OVERLAP", "200"))
EMBEDDING_BATCH_SIZE: int = int(os.environ.get("EMBEDDING_BATCH_SIZE", "64"))

EXTENSIONES_SOPORTADAS: set[str] = {".pdf", ".docx", ".xlsx", ".xls"}

# Las 7 colecciones del corpus de la CGR
COLECCIONES: list[str] = [
    "normativo",
    "institucional",
    "academico",
    "tecnico_tic",
    "estadistico",
    "jurisprudencial",
    "auditoria",
]

# Modo de chunking segun coleccion
MODO_CHUNKING: dict[str, str] = {
    "normativo": "juridico",
    "jurisprudencial": "juridico",
    "institucional": "institucional",
    "auditoria": "institucional",
    "academico": "general",
    "tecnico_tic": "general",
    "estadistico": "general",
}


# ---------------------------------------------------------------------------
# Estadisticas
# ---------------------------------------------------------------------------

@dataclass
class EstadisticasIngestion:
    archivos_procesados: int = 0
    archivos_omitidos: int = 0
    fragmentos_generados: int = 0
    embeddings_almacenados: int = 0
    errores: int = 0
    tiempo_inicio: float = field(default_factory=time.time)

    def resumen(self) -> str:
        duracion = time.time() - self.tiempo_inicio
        return (
            "\n" + "=" * 60 + "\n"
            "  RESUMEN DE INGESTION DEL CORPUS\n"
            + "=" * 60 + "\n"
            f"  Archivos procesados   : {self.archivos_procesados}\n"
            f"  Archivos omitidos     : {self.archivos_omitidos}\n"
            f"  Fragmentos generados  : {self.fragmentos_generados}\n"
            f"  Embeddings almacenados: {self.embeddings_almacenados}\n"
            f"  Errores               : {self.errores}\n"
            f"  Duracion              : {duracion:.1f} s\n"
            + "=" * 60
        )


# ---------------------------------------------------------------------------
# Extraccion de texto
# ---------------------------------------------------------------------------

def extraer_texto_pdf(ruta: Path) -> list[tuple[int | None, str]]:
    """Extrae texto pagina a pagina de un PDF."""
    try:
        import fitz
    except ImportError:
        logger.error("PyMuPDF no instalado. pip install pymupdf")
        return []

    paginas: list[tuple[int | None, str]] = []
    with fitz.open(str(ruta)) as doc:
        for i, pagina in enumerate(doc, start=1):
            texto = pagina.get_text("text")
            if texto.strip():
                paginas.append((i, texto))
    return paginas


def extraer_texto_docx(ruta: Path) -> list[tuple[int | None, str]]:
    """Extrae texto de un archivo DOCX."""
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx no instalado. pip install python-docx")
        return []

    doc = Document(str(ruta))
    parrafos = []
    for p in doc.paragraphs:
        if p.text.strip():
            parrafos.append(p.text)
    for tabla in doc.tables:
        for fila in tabla.rows:
            celdas = [c.text.strip() for c in fila.cells if c.text.strip()]
            if celdas:
                parrafos.append(" | ".join(celdas))
    texto = "\n\n".join(parrafos)
    return [(None, texto)] if texto else []


def extraer_texto_xlsx(ruta: Path) -> list[tuple[int | None, str]]:
    """Extrae contenido de un XLSX."""
    try:
        import openpyxl
    except ImportError:
        logger.error("openpyxl no instalado. pip install openpyxl")
        return []

    wb = openpyxl.load_workbook(str(ruta), read_only=True, data_only=True)
    textos: list[str] = []
    for hoja in wb.sheetnames:
        ws = wb[hoja]
        filas = []
        for fila in ws.iter_rows(values_only=True):
            celdas = [str(c) for c in fila if c is not None]
            if celdas:
                filas.append(" | ".join(celdas))
        if filas:
            textos.append(f"[Hoja: {hoja}]\n" + "\n".join(filas))
    wb.close()
    contenido = "\n\n".join(textos)
    return [(None, contenido)] if contenido else []


EXTRACTORES = {
    ".pdf": extraer_texto_pdf,
    ".docx": extraer_texto_docx,
    ".xlsx": extraer_texto_xlsx,
    ".xls": extraer_texto_xlsx,
}


# ---------------------------------------------------------------------------
# Fragmentacion con modo segun coleccion
# ---------------------------------------------------------------------------

def fragmentar_texto(
    texto: str,
    coleccion: str,
    tamano: int = CHUNK_SIZE,
    solapamiento: int = CHUNK_OVERLAP,
) -> list[str]:
    """Divide texto en fragmentos. Usa modo juridico para normativo/jurisprudencial."""
    # Importar chunking del modulo RAG si esta disponible
    try:
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "backend"))
        from app.rag.chunking import dividir_en_fragmentos, detectar_modo_por_coleccion

        modo = detectar_modo_por_coleccion(coleccion)
        fragmentos = dividir_en_fragmentos(texto, tamano=tamano, solapamiento=solapamiento, modo=modo)
        return [f.contenido for f in fragmentos]
    except ImportError:
        pass

    # Fallback simple si no se puede importar el modulo
    if len(texto) <= tamano:
        return [texto] if texto.strip() else []

    fragmentos: list[str] = []
    inicio = 0
    while inicio < len(texto):
        fin = inicio + tamano
        fragmento = texto[inicio:fin].strip()
        if fragmento:
            fragmentos.append(fragmento)
        inicio += tamano - solapamiento
    return fragmentos


# ---------------------------------------------------------------------------
# Generacion de embeddings
# ---------------------------------------------------------------------------

def generar_embeddings_batch(textos: list[str]) -> list[list[float]]:
    """Genera embeddings usando la API configurada en las variables de entorno."""
    # Intentar importar el modulo RAG
    try:
        sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "backend"))
        from app.rag.embeddings import generar_embeddings
        return generar_embeddings(textos)
    except ImportError:
        pass

    # Fallback: usar openai directamente
    try:
        from openai import OpenAI

        base_url = os.environ.get("EMBEDDINGS_BASE_URL", "http://localhost:11434/v1")
        api_key = os.environ.get("EMBEDDINGS_API_KEY", "ollama")
        modelo = os.environ.get("EMBEDDINGS_MODEL", "nomic-embed-text")

        cliente = OpenAI(base_url=base_url, api_key=api_key)
        respuesta = cliente.embeddings.create(input=textos, model=modelo)
        return [item.embedding for item in respuesta.data]
    except Exception as exc:
        logger.error("Error al generar embeddings: %s", exc)
        # Vector cero como fallback (permite continuar sin API)
        dim = int(os.environ.get("EMBEDDINGS_DIMENSIONES", "768"))
        logger.warning("Usando vectores cero de dimension %d como fallback.", dim)
        return [[0.0] * dim for _ in textos]


# ---------------------------------------------------------------------------
# Almacenamiento en pgvector
# ---------------------------------------------------------------------------

def registrar_documento(session, ruta: Path, coleccion: str) -> str:
    """Registra un documento en la tabla documentos y retorna su ID."""
    from sqlalchemy import text as sql_text

    doc_id = str(uuid.uuid4())
    tamano = ruta.stat().st_size
    tipo_mime = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ".xls": "application/vnd.ms-excel",
    }.get(ruta.suffix.lower(), "application/octet-stream")

    # Hash del contenido para deteccion de duplicados
    hash_contenido = hashlib.sha256(ruta.read_bytes()).hexdigest()

    # Verificar si ya existe por hash
    existente = session.execute(
        sql_text("SELECT id FROM documentos WHERE hash_contenido = :hash"),
        {"hash": hash_contenido},
    ).fetchone()

    if existente:
        logger.info("  Documento ya existe (hash=%s...), se omite.", hash_contenido[:12])
        return ""

    session.execute(
        sql_text("""
            INSERT INTO documentos (id, nombre_archivo, tipo_mime, tamano_bytes, coleccion,
                                    estado, ruta_almacenamiento, hash_contenido)
            VALUES (:id, :nombre, :tipo, :tamano, :col, 'procesando', :ruta, :hash)
        """),
        {
            "id": doc_id,
            "nombre": ruta.name,
            "tipo": tipo_mime,
            "tamano": tamano,
            "col": coleccion,
            "ruta": str(ruta),
            "hash": hash_contenido,
        },
    )
    return doc_id


def almacenar_fragmentos(
    session,
    documento_id: str,
    coleccion: str,
    fragmentos_texto: list[str],
    paginas: list[int | None],
    embeddings: list[list[float]],
) -> int:
    """Inserta fragmentos con embeddings en fragmentos_vectoriales."""
    from sqlalchemy import text as sql_text

    insertados = 0
    for i, (texto, pagina, emb) in enumerate(zip(fragmentos_texto, paginas, embeddings)):
        frag_id = str(uuid.uuid4())
        emb_str = "[" + ",".join(str(v) for v in emb) + "]"

        try:
            session.execute(
                sql_text("""
                    INSERT INTO fragmentos_vectoriales
                        (id, documento_id, contenido, coleccion, pagina, posicion_fragmento, embedding)
                    VALUES
                        (:id, :doc_id, :contenido, :col, :pagina, :pos, CAST(:emb AS vector))
                """),
                {
                    "id": frag_id,
                    "doc_id": documento_id,
                    "contenido": texto,
                    "col": coleccion,
                    "pagina": pagina,
                    "pos": i,
                    "emb": emb_str,
                },
            )
            insertados += 1
        except Exception as exc:
            logger.error("Error al insertar fragmento %d: %s", i, exc)
            session.rollback()

    return insertados


def actualizar_documento_completado(session, documento_id: str, total_fragmentos: int) -> None:
    """Marca el documento como indexado y actualiza el conteo de fragmentos."""
    from sqlalchemy import text as sql_text

    session.execute(
        sql_text("""
            UPDATE documentos
            SET estado = 'indexado', total_fragmentos = :total
            WHERE id = :id
        """),
        {"id": documento_id, "total": total_fragmentos},
    )


# ---------------------------------------------------------------------------
# Escaneo del corpus
# ---------------------------------------------------------------------------

def escanear_corpus(base: Path) -> Iterator[tuple[str, Path]]:
    """Recorre las 7 colecciones y genera (coleccion, ruta_archivo)."""
    for coleccion in COLECCIONES:
        directorio = base / coleccion
        if not directorio.is_dir():
            logger.info("Directorio no encontrado: %s — se omite.", directorio)
            continue

        archivos = sorted(directorio.rglob("*"))
        logger.info("Coleccion '%s': %d archivos encontrados.", coleccion,
                     sum(1 for a in archivos if a.is_file() and a.suffix.lower() in EXTENSIONES_SOPORTADAS))

        for archivo in archivos:
            if archivo.is_file() and archivo.suffix.lower() in EXTENSIONES_SOPORTADAS:
                yield coleccion, archivo


# ---------------------------------------------------------------------------
# Proceso principal
# ---------------------------------------------------------------------------

def ingestar() -> None:
    """Ejecuta el flujo completo de ingestion del corpus."""
    if not DATABASE_URL:
        logger.error("DATABASE_URL o DATABASE_URL_SYNC no configurada.")
        sys.exit(1)

    from sqlalchemy import create_engine
    from sqlalchemy.orm import sessionmaker

    engine = create_engine(DATABASE_URL, echo=False, pool_pre_ping=True)
    Session = sessionmaker(bind=engine)
    session = Session()
    stats = EstadisticasIngestion()

    # Verificar directorio corpus
    if not CORPUS_DIR.is_dir():
        logger.error("Directorio de corpus no encontrado: %s", CORPUS_DIR)
        logger.info("Ejecute primero: scripts/copiar_corpus.bat (Windows) o scripts/copiar_corpus.sh (Linux)")
        sys.exit(1)

    logger.info("=" * 60)
    logger.info("  INGESTION DEL CORPUS — CecilIA v2")
    logger.info("=" * 60)
    logger.info("Corpus: %s", CORPUS_DIR.resolve())
    logger.info("DB: %s", DATABASE_URL[:50] + "...")
    logger.info("")

    for coleccion, ruta in escanear_corpus(CORPUS_DIR):
        logger.info("[%s] Procesando: %s", coleccion, ruta.name)
        extractor = EXTRACTORES.get(ruta.suffix.lower())

        if extractor is None:
            stats.archivos_omitidos += 1
            continue

        try:
            paginas_texto = extractor(ruta)
        except Exception as exc:
            logger.error("  Error al extraer texto: %s", exc)
            stats.errores += 1
            continue

        if not paginas_texto:
            logger.info("  Sin contenido extraible, se omite.")
            stats.archivos_omitidos += 1
            continue

        # Registrar documento en la BD
        try:
            documento_id = registrar_documento(session, ruta, coleccion)
            if not documento_id:
                stats.archivos_omitidos += 1
                session.rollback()
                continue
        except Exception as exc:
            logger.error("  Error al registrar documento: %s", exc)
            session.rollback()
            stats.errores += 1
            continue

        # Fragmentar cada pagina
        todos_fragmentos: list[str] = []
        todas_paginas: list[int | None] = []

        for pagina_num, texto in paginas_texto:
            trozos = fragmentar_texto(texto, coleccion)
            for trozo in trozos:
                todos_fragmentos.append(trozo)
                todas_paginas.append(pagina_num)

        if not todos_fragmentos:
            logger.info("  Sin fragmentos generables, se omite.")
            session.rollback()
            stats.archivos_omitidos += 1
            continue

        # Generar embeddings y almacenar por lotes
        total_insertados = 0
        for i in range(0, len(todos_fragmentos), EMBEDDING_BATCH_SIZE):
            lote_textos = todos_fragmentos[i: i + EMBEDDING_BATCH_SIZE]
            lote_paginas = todas_paginas[i: i + EMBEDDING_BATCH_SIZE]

            try:
                embs = generar_embeddings_batch(lote_textos)
            except Exception as exc:
                logger.error("  Error embeddings lote %d: %s", i, exc)
                stats.errores += 1
                continue

            insertados = almacenar_fragmentos(
                session, documento_id, coleccion,
                lote_textos, lote_paginas, embs,
            )
            total_insertados += insertados

        # Actualizar estado del documento y hacer commit por documento
        try:
            actualizar_documento_completado(session, documento_id, total_insertados)
            session.commit()
        except Exception as exc:
            logger.error("  Error al commit documento %s: %s", ruta.name, exc)
            session.rollback()
            stats.errores += 1
            continue

        stats.archivos_procesados += 1
        stats.fragmentos_generados += len(todos_fragmentos)
        stats.embeddings_almacenados += total_insertados

        logger.info(
            "  -> %d fragmentos, %d embeddings almacenados",
            len(todos_fragmentos), total_insertados,
        )

    # Crear indice IVFFlat si hay suficientes datos
    try:
        from sqlalchemy import text as sql_text
        conteo = session.execute(
            sql_text("SELECT COUNT(*) FROM fragmentos_vectoriales")
        ).scalar()

        if conteo and conteo >= 100:
            logger.info("Creando indice IVFFlat (%d fragmentos)...", conteo)
            listas = min(max(conteo // 10, 10), 1000)
            session.execute(sql_text(
                f"CREATE INDEX IF NOT EXISTS ix_fv_embedding_ivfflat "
                f"ON fragmentos_vectoriales "
                f"USING ivfflat (embedding vector_cosine_ops) "
                f"WITH (lists = {listas})"
            ))
            session.commit()
            logger.info("Indice IVFFlat creado con %d listas.", listas)
    except Exception as exc:
        logger.warning("No se pudo crear indice IVFFlat: %s", exc)

    session.close()
    logger.info(stats.resumen())


if __name__ == "__main__":
    ingestar()
