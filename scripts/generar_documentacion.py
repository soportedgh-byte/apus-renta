#!/usr/bin/env python3
"""
CecilIA v2 — Generador de Documentacion Tecnica Institucional
Contraloria General de la Republica de Colombia
Contraloria Delegada para el Sector TIC — CD-TIC

Genera 17 documentos DOCX + 3 XLSX + 3 documentos adicionales
con formato institucional CGR profesional.

Autor: Equipo Tecnico CecilIA — CD-TIC-CGR
Fecha: Abril 2026
"""

from __future__ import annotations

import io
import os
import sys
import time
from datetime import datetime
from pathlib import Path
from typing import Any

# Asegurar que el backend este en el path
RAIZ = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(RAIZ / "backend"))

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Mm, Pt, RGBColor, Emu

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    TIENE_OPENPYXL = True
except ImportError:
    TIENE_OPENPYXL = False
    print("[AVISO] openpyxl no disponible — XLSX no se generaran")

# ── Constantes de estilo ────────────────────────────────────────────────────

COLOR_AZUL = RGBColor(0x1A, 0x52, 0x76)
COLOR_AZUL_CLARO = RGBColor(0xEB, 0xF5, 0xFB)
COLOR_ORO = RGBColor(0xC9, 0xA8, 0x4C)
COLOR_ROJO = RGBColor(0xC0, 0x39, 0x2B)
COLOR_BLANCO = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_NEGRO = RGBColor(0x00, 0x00, 0x00)
COLOR_GRIS = RGBColor(0x5F, 0x63, 0x68)
FUENTE_TITULO = "Georgia"
FUENTE_CUERPO = "Calibri"
FUENTE_CODIGO = "Consolas"

CARPETA_SALIDA = RAIZ / "docs" / "entregables"
LOGO_CGR = RAIZ / "frontend" / "public" / "logo-cgr.png"
LOGO_CECILIA = RAIZ / "frontend" / "public" / "logo-cecilia.png"

EQUIPO = [
    ("Dr. Omar Javier Contreras Socarras", "Impulsor del proyecto", "Contralor Delegado TIC (Concepcion)"),
    ("Dr. Hector Hernan Gonzalez Naranjo", "Contralor Delegado TIC", "Direccion operativa"),
    ("Dr. Juan Carlos Cobo Gomez", "Director DES", "Control fiscal macro"),
    ("Dr. Jose Fernando Ramirez Munoz", "Director DVF", "Control fiscal micro"),
    ("Dra. Karen Tatiana Suevis Gomez", "Coordinadora Juridica", "Marco legal y etica"),
    ("Ing. Gustavo Adolfo Castillo Romero", "Lider Tecnico", "Arquitectura y desarrollo"),
]


# ══════════════════════════════════════════════════════════════════════════════
# CLASE BASE: Documento DOCX institucional
# ══════════════════════════════════════════════════════════════════════════════

class DocCGR:
    """Clase base para generar documentos DOCX con formato institucional CGR."""

    def __init__(self, titulo: str, subtitulo: str = "", nombre_archivo: str = "documento"):
        self.titulo = titulo
        self.subtitulo = subtitulo
        self.nombre_archivo = nombre_archivo
        self.doc = Document()
        self._configurar_estilos()
        self._configurar_margenes()

    def _configurar_margenes(self):
        for seccion in self.doc.sections:
            seccion.top_margin = Cm(2.5)
            seccion.bottom_margin = Cm(2.5)
            seccion.left_margin = Cm(2.0)
            seccion.right_margin = Cm(2.0)
            seccion.page_width = Mm(210)
            seccion.page_height = Mm(297)

    def _configurar_estilos(self):
        estilo = self.doc.styles["Normal"]
        estilo.font.name = FUENTE_CUERPO
        estilo.font.size = Pt(11)
        estilo.font.color.rgb = COLOR_NEGRO
        estilo.paragraph_format.space_before = Pt(2)
        estilo.paragraph_format.space_after = Pt(4)
        estilo.paragraph_format.line_spacing = 1.15

    def agregar_logo(self, ancho=Cm(6)):
        if LOGO_CGR.exists():
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run()
            run.add_picture(str(LOGO_CGR), width=ancho)

    def agregar_logo_cecilia(self, ancho=Cm(2.5)):
        if LOGO_CECILIA.exists():
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run()
            run.add_picture(str(LOGO_CECILIA), width=ancho)

    def agregar_linea_dorada(self):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run("_" * 80)
        run.font.color.rgb = COLOR_ORO
        run.font.size = Pt(6)

    def agregar_portada(self):
        """Portada institucional completa."""
        self.agregar_logo(ancho=Cm(8))
        self.agregar_logo_cecilia(ancho=Cm(3))

        # Titulo
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(20)
        run = p.add_run(self.titulo.upper())
        run.font.name = FUENTE_TITULO
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = COLOR_AZUL

        # Subtitulo
        if self.subtitulo:
            p2 = self.doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.paragraph_format.space_before = Pt(6)
            run2 = p2.add_run(self.subtitulo)
            run2.font.name = FUENTE_CUERPO
            run2.font.size = Pt(14)
            run2.font.color.rgb = COLOR_GRIS

        self.agregar_linea_dorada()

        # Datos institucionales
        lineas_inst = [
            ("Contraloria General de la Republica", True, Pt(12)),
            ("Contraloria Delegada para el Sector TIC — CD-TIC", False, Pt(10)),
            ("", False, Pt(8)),
            ("Proyecto concebido e impulsado por el", False, Pt(9)),
            ("Dr. Omar Javier Contreras Socarras", True, Pt(10)),
            ("Contralor Delegado TIC", False, Pt(9)),
            ("", False, Pt(12)),
            ("Version 1.0 — Abril 2026", False, Pt(10)),
            ("DOCUMENTO DE USO INTERNO — EQUIPO TECNICO PILOTO CECILIA", False, Pt(8)),
        ]

        for texto, negrita, tamano in lineas_inst:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            if texto:
                run = p.add_run(texto)
                run.font.name = FUENTE_CUERPO
                run.font.size = tamano
                run.font.bold = negrita
                if "Version" in texto:
                    run.font.color.rgb = COLOR_ORO
                elif "DOCUMENTO" in texto:
                    run.font.color.rgb = COLOR_GRIS
                    run.font.italic = True

        self.doc.add_page_break()

    def agregar_encabezado_pagina(self):
        """Encabezado institucional en cada pagina."""
        self.agregar_logo(ancho=Cm(5))

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("CONTRALORIA GENERAL DE LA REPUBLICA")
        run.font.name = FUENTE_TITULO
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = COLOR_AZUL

        p2 = self.doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(2)
        run2 = p2.add_run("Contraloria Delegada para el Sector TIC — CD-TIC")
        run2.font.name = FUENTE_CUERPO
        run2.font.size = Pt(11)

        p3 = self.doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_after = Pt(4)
        run3 = p3.add_run("Piloto de Inteligencia Artificial CecilIA")
        run3.font.name = FUENTE_CUERPO
        run3.font.size = Pt(10)
        run3.font.italic = True
        run3.font.color.rgb = COLOR_GRIS

        self.agregar_linea_dorada()

    def h1(self, texto: str, numeracion: str = ""):
        """Heading 1: seccion principal."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        contenido = f"{numeracion} {texto}".strip() if numeracion else texto
        run = p.add_run(contenido)
        run.font.name = FUENTE_TITULO
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = COLOR_AZUL

    def h2(self, texto: str, numeracion: str = ""):
        """Heading 2: subseccion."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        contenido = f"{numeracion} {texto}".strip() if numeracion else texto
        run = p.add_run(contenido)
        run.font.name = FUENTE_TITULO
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x0A, 0x0F, 0x1A)

    def h3(self, texto: str, numeracion: str = ""):
        """Heading 3: sub-subseccion."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        contenido = f"{numeracion} {texto}".strip() if numeracion else texto
        run = p.add_run(contenido)
        run.font.name = FUENTE_CUERPO
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = COLOR_AZUL

    def parrafo(self, texto: str):
        """Parrafo justificado."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(texto)
        run.font.name = FUENTE_CUERPO
        run.font.size = Pt(11)

    def lista(self, items: list[str], ordenada: bool = False):
        """Lista con guiones o numeros."""
        for i, item in enumerate(items, 1):
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            prefijo = f"{i}. " if ordenada else "- "
            run = p.add_run(f"{prefijo}{item}")
            run.font.name = FUENTE_CUERPO
            run.font.size = Pt(11)

    def nota(self, texto: str):
        """Nota destacada con fondo amarillo."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(f"NOTA: {texto}")
        run.font.name = FUENTE_CUERPO
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0xF3, 0x9C, 0x12)

    def codigo(self, texto: str):
        """Bloque de codigo."""
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(texto)
        run.font.name = FUENTE_CODIGO
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_GRIS

    def tabla(self, encabezados: list[str], filas: list[list[str]]):
        """Tabla profesional con formato CGR."""
        tabla = self.doc.add_table(rows=1 + len(filas), cols=len(encabezados))
        tabla.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Encabezados
        for j, enc in enumerate(encabezados):
            celda = tabla.rows[0].cells[j]
            celda.text = ""
            p = celda.paragraphs[0]
            run = p.add_run(enc)
            run.font.name = FUENTE_CUERPO
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = COLOR_BLANCO
            # Fondo azul
            tc = celda._tc
            tcPr = tc.get_or_add_tcPr()
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "1A5276")
            shading.set(qn("w:val"), "clear")
            tcPr.append(shading)

        # Filas
        for i, fila in enumerate(filas):
            for j, valor in enumerate(fila):
                celda = tabla.rows[i + 1].cells[j]
                celda.text = ""
                p = celda.paragraphs[0]
                run = p.add_run(str(valor))
                run.font.name = FUENTE_CUERPO
                run.font.size = Pt(9)
                # Filas alternas
                if i % 2 == 0:
                    tc = celda._tc
                    tcPr = tc.get_or_add_tcPr()
                    shading = OxmlElement("w:shd")
                    shading.set(qn("w:fill"), "EBF5FB")
                    shading.set(qn("w:val"), "clear")
                    tcPr.append(shading)

        self.doc.add_paragraph()  # Espacio despues de tabla

    def agregar_contraportada(self):
        """Contraportada con equipo tecnico."""
        self.doc.add_page_break()

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(60)
        run = p.add_run("Equipo Tecnico Piloto CecilIA")
        run.font.name = FUENTE_TITULO
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = COLOR_AZUL

        self.agregar_linea_dorada()

        self.tabla(
            ["Nombre", "Cargo", "Rol en el proyecto"],
            [[n, c, r] for n, c, r in EQUIPO],
        )

        for texto, negrita in [
            ("Contraloria Delegada para el Sector TIC — CD-TIC-CGR", True),
            ("Bogota D.C., Colombia — 2026", False),
        ]:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            run = p.add_run(texto)
            run.font.name = FUENTE_CUERPO
            run.font.size = Pt(11)
            run.font.bold = negrita

    def agregar_nota_ia(self):
        """Disclaimer obligatorio Circular 023."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        run = p.add_run(
            "Este documento fue generado con asistencia de inteligencia artificial (CecilIA v2). "
            "Todos los datos, analisis y conclusiones deben ser verificados y validados por el equipo "
            "auditor antes de su uso oficial, conforme a lo establecido en la Circular 023 de la CGR."
        )
        run.font.name = FUENTE_CUERPO
        run.font.size = Pt(8)
        run.font.italic = True
        run.font.color.rgb = COLOR_GRIS

    def guardar(self) -> Path:
        """Guarda el documento y retorna la ruta."""
        ruta = CARPETA_SALIDA / f"{self.nombre_archivo}.docx"
        self.doc.save(str(ruta))
        return ruta

    def contar_paginas_aprox(self) -> int:
        """Estima el numero de paginas basado en parrafos."""
        return max(1, len(self.doc.paragraphs) // 25)


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENTO 1: Manual de Usuario
# ══════════════════════════════════════════════════════════════════════════════

def generar_manual_usuario() -> Path:
    doc = DocCGR("Manual de Usuario de CecilIA v2",
                 "Guia completa para auditores y funcionarios de la CGR",
                 "01_Manual_Usuario_CecilIA_v2")
    doc.agregar_portada()
    doc.agregar_encabezado_pagina()

    # 1. Introduccion
    doc.h1("Introduccion", "1.")
    doc.h2("Que es CecilIA?", "1.1")
    doc.parrafo(
        "CecilIA v2 es el Sistema de Inteligencia Artificial para Control Fiscal de la Contraloria General "
        "de la Republica de Colombia (CGR), desarrollado por la Contraloria Delegada para el Sector TIC (CD-TIC). "
        "El nombre CecilIA deriva del latin caecus — 'aquel que no se deslumbra por lo superficial' — "
        "simbolizando objetividad y analisis basado en datos."
    )
    doc.parrafo(
        "CecilIA asiste a auditores de dos direcciones: DES (Direccion de Estudios Sectoriales) para control "
        "fiscal macro, y DVF (Direccion de Vigilancia Fiscal) para control fiscal micro. CecilIA es un asistente, "
        "nunca un reemplazo del auditor. Genera borradores y sugerencias que el ser humano revisa, modifica y valida."
    )

    doc.h2("Para quien es este manual?", "1.2")
    doc.parrafo("Este manual esta dirigido a todos los usuarios del sistema CecilIA v2:")
    doc.lista([
        "Auditores DES y DVF: uso diario del chat, formatos y hallazgos",
        "Directores: supervision, aprobacion y analitica",
        "Coordinadores: gestion de politicas y auditorias",
        "Aprendices: modulo de capacitacion y tutoria",
        "Administradores: configuracion del sistema",
    ])

    doc.h2("Requisitos de acceso", "1.3")
    doc.tabla(
        ["Requisito", "Detalle"],
        [
            ["Navegador", "Chrome 90+, Firefox 88+, Edge 90+ (actualizados)"],
            ["Resolucion", "Minimo 1280x720 — Recomendado 1920x1080"],
            ["Red", "Conexion a la red interna CGR o VPN autorizada"],
            ["Credenciales", "Usuario y contrasena proporcionados por el administrador"],
        ],
    )

    # 2. Acceso al sistema
    doc.h1("Acceso al sistema", "2.")
    doc.h2("Como iniciar sesion", "2.1")
    doc.parrafo(
        "Ingrese a la URL del sistema (proporcionada por el area de TIC). En la pantalla de login, "
        "ingrese su nombre de usuario y contrasena. Haga clic en 'Iniciar sesion'."
    )
    doc.parrafo(
        "La pantalla de login muestra los logos institucionales de la CGR y CecilIA, asi como la mencion "
        "al Dr. Omar Javier Contreras Socarras, impulsor del proyecto."
    )

    doc.h2("Seleccion de direccion de trabajo (DES/DVF)", "2.2")
    doc.parrafo(
        "Despues de iniciar sesion, se le presentara una pantalla para seleccionar su direccion de trabajo. "
        "Si su rol solo tiene acceso a una direccion, sera redirigido automaticamente. "
        "Los usuarios con acceso dual (admin, lider tecnico, coordinador) veran ambas opciones."
    )

    doc.h2("Descripcion de la interfaz principal", "2.3")
    doc.parrafo(
        "La interfaz principal consta de: Sidebar izquierdo con navegacion a todos los modulos "
        "(Chat IA, Capacitacion, Workspace, Auditorias, Hallazgos, Formatos, Observatorio, Analitica); "
        "Header superior con estado de conexion, modelo LLM activo y datos del usuario; "
        "Area central de contenido que cambia segun el modulo seleccionado; "
        "Footer con informacion legal y creditos del proyecto."
    )

    doc.h2("Declaracion de responsabilidad — Circular 023", "2.4")
    doc.parrafo(
        "Al iniciar sesion por primera vez, se mostrara un modal con la declaracion de responsabilidad "
        "conforme a la Circular 023 de la CGR. El usuario debe leer y aceptar esta declaracion, "
        "que establece que CecilIA es un asistente y toda respuesta debe ser validada por el auditor."
    )

    # 3. Chat con CecilIA
    doc.h1("Modulo de Chat con CecilIA", "3.")
    doc.h2("Como hacer preguntas efectivas", "3.1")
    doc.parrafo(
        "CecilIA funciona mejor con preguntas claras y especificas. Incluya contexto relevante como "
        "la entidad auditada, la vigencia, y el tipo de auditoria. Ejemplos de buenas preguntas:"
    )
    doc.lista([
        "Que normatividad aplica para auditar la contratacion publica de una entidad territorial?",
        "Genera un borrador de hallazgo administrativo sobre falta de publicacion en SECOP",
        "Cual es la estructura del Formato F14 de observaciones?",
        "Explica el concepto de materialidad segun la NIA 320",
    ])

    doc.h2("Acciones rapidas segun su rol", "3.2")
    doc.parrafo("El chat presenta acciones rapidas personalizadas segun la direccion activa:")
    doc.tabla(
        ["Direccion", "Acciones rapidas"],
        [
            ["DVF", "Pre-planeacion, Materialidad, Hallazgo, Formato CGR"],
            ["DES", "Analisis presupuestal, Regalias, Politica publica, Alertas sectoriales"],
        ],
    )

    doc.h2("Como interpretar las fuentes consultadas", "3.3")
    doc.parrafo(
        "Cada respuesta de CecilIA incluye las fuentes consultadas del RAG (Retrieval-Augmented Generation). "
        "Las fuentes se muestran al final de la respuesta con el nombre del documento, la coleccion y "
        "un indicador de relevancia. Esto permite al auditor verificar la informacion."
    )

    doc.h2("Gestion de conversaciones", "3.4")
    doc.parrafo(
        "El panel izquierdo del chat muestra el historial de conversaciones. Puede crear nuevas conversaciones, "
        "renombrar las existentes haciendo clic en el icono de edicion, o eliminarlas con el icono de papelera."
    )

    # 4. Formatos CGR
    doc.h1("Modulo de Formatos CGR (solo DVF)", "4.")
    doc.h2("Formatos disponibles por fase", "4.1")
    doc.parrafo(
        "CecilIA puede generar 30 formatos de la Guia de Auditoria Financiera (GAF) de la CGR, "
        "organizados por las 5 fases del proceso auditor."
    )
    doc.tabla(
        ["Fase", "Formatos", "Cantidad"],
        [
            ["Pre-planeacion", "F01 a F06", "6"],
            ["Planeacion", "F07 a F13", "7"],
            ["Ejecucion", "F14 a F22", "9"],
            ["Informe", "F23 a F27", "5"],
            ["Seguimiento", "F28 a F30", "3"],
        ],
    )

    doc.h2("Como generar un formato", "4.2")
    doc.parrafo(
        "Navegue al modulo 'Formatos' en el sidebar. Seleccione el formato deseado de la lista organizada "
        "por fases. Complete los campos requeridos (entidad, vigencia, datos de auditoria) y haga clic en "
        "'Generar'. El sistema creara un documento DOCX con encabezado institucional CGR."
    )

    doc.h2("Descarga y verificacion", "4.3")
    doc.parrafo(
        "Antes de descargar cualquier formato, debe marcar el checkbox 'He revisado y valido este documento' "
        "conforme a la Circular 023. El archivo DOCX incluye pie de pagina con disclaimer de IA."
    )

    # 5. Hallazgos
    doc.h1("Modulo de Hallazgos (solo DVF)", "5.")
    doc.h2("Como crear un hallazgo", "5.1")
    doc.parrafo(
        "Navegue a 'Hallazgos' en el sidebar y haga clic en 'Nuevo hallazgo'. Complete los campos "
        "obligatorios incluyendo los 4 elementos del hallazgo."
    )

    doc.h2("Los 4 elementos obligatorios", "5.2")
    doc.tabla(
        ["Elemento", "Descripcion"],
        [
            ["Condicion", "Lo que el auditor encontro (situacion actual)"],
            ["Criterio", "La norma o regla contra la cual se evalua"],
            ["Causa", "Por que se produjo la diferencia"],
            ["Efecto", "Consecuencia o impacto de la irregularidad"],
        ],
    )

    doc.h2("Connotaciones", "5.3")
    doc.lista([
        "Administrativa: Incumplimiento de funciones o deberes",
        "Fiscal: Presunto dano patrimonial al Estado",
        "Disciplinaria: Falta disciplinaria del servidor publico",
        "Penal: Presunta conducta delictiva",
    ])

    doc.h2("Workflow de aprobacion", "5.4")
    doc.parrafo(
        "Los hallazgos siguen un flujo: Borrador → En revision → Aprobado. "
        "Solo los directores pueden aprobar hallazgos para su inclusion en el informe final."
    )

    # 6. Capacitacion
    doc.h1("Modulo de Capacitacion", "6.")
    doc.h2("Rutas de aprendizaje", "6.1")
    doc.parrafo("CecilIA ofrece 4 rutas de aprendizaje adaptativas:")
    doc.lista([
        "Conoce la CGR: Historia, estructura y mision de la Contraloria",
        "Auditoria DVF: Paso a paso del proceso auditor micro",
        "Estudios Sectoriales DES: Metodologia de control fiscal macro",
        "Normativa Esencial: Leyes, decretos y circular 023",
    ])

    doc.h2("Sistema de aprendizaje adaptativo", "6.2")
    doc.parrafo(
        "Al ingresar al modulo, puede completar un cuestionario VARK para determinar su estilo de "
        "aprendizaje (lector, auditivo, visual o kinestesico). CecilIA adaptara el contenido de las "
        "lecciones segun su estilo predominante."
    )

    doc.h2("Gamificacion", "6.3")
    doc.parrafo(
        "El sistema incluye gamificacion con puntos XP, niveles (Practicante, Asistente, Analista, Auditor, Experto), "
        "rachas de estudio diario e insignias por logros especificos."
    )

    doc.h2("Biblioteca de recursos", "6.4")
    doc.parrafo(
        "La biblioteca genera contenido educativo con IA: podcasts con 2 voces colombianas, "
        "flashcards con taxonomia de Bloom, infografias renderizadas y manuales DOCX con logos CGR."
    )

    # 7. Observatorio TIC
    doc.h1("Modulo de Observatorio TIC (solo DES)", "7.")
    doc.parrafo(
        "El Observatorio TIC monitorea automaticamente fuentes del sector TIC (MinTIC, CRC, ANE) "
        "y clasifica alertas por nivel de relevancia usando IA. Los directores DES pueden iniciar "
        "analisis desde cualquier alerta detectada."
    )

    # 8. Analitica
    doc.h1("Modulo de Analitica (solo Directores)", "8.")
    doc.parrafo(
        "Dashboard ejecutivo con metricas de uso del sistema: total de consultas, formatos generados, "
        "hallazgos creados, tiempo promedio de respuesta y tendencias mensuales. "
        "Permite exportar informes periodicos."
    )

    # 9. Workspace
    doc.h1("Workspace", "9.")
    doc.parrafo(
        "El Workspace permite cargar documentos para analisis. Tipos soportados: PDF, DOCX, XLSX, CSV, TXT. "
        "Los documentos cargados se procesan y quedan disponibles para consulta en el chat."
    )

    # 10. Perfil
    doc.h1("Mi Perfil y Configuracion", "10.")
    doc.parrafo(
        "Desde el menu de perfil puede ver sus datos, cambiar contrasena y configurar preferencias "
        "del sistema como la direccion activa por defecto."
    )

    # 11. FAQ
    doc.h1("Preguntas Frecuentes", "11.")
    faqs = [
        ("CecilIA reemplaza al auditor?", "No. CecilIA es un asistente que genera borradores y sugerencias. Toda respuesta debe ser validada por el auditor conforme a la Circular 023."),
        ("Puedo confiar en las respuestas de CecilIA?", "CecilIA consulta fuentes normativas reales pero puede cometer errores. Siempre verifique contra la norma original."),
        ("Mis datos estan seguros?", "Si. Los datos se procesan en servidores institucionales con cifrado AES-256 y TLS 1.3."),
        ("Que hago si CecilIA da una respuesta incorrecta?", "Use los botones de feedback (pulgar arriba/abajo) y reporte al equipo tecnico."),
        ("Puedo usar CecilIA para generar informes finales?", "No. La Circular 023 limita el uso de IA a borradores y sugerencias, nunca como producto final."),
        ("Como actualizo mis credenciales?", "Contacte al administrador del sistema o use la opcion de cambio de contrasena en su perfil."),
        ("CecilIA funciona sin internet?", "En la red interna CGR si. El Desktop Agent (futuro) permitira uso local."),
        ("Puedo adjuntar documentos al chat?", "Si, a traves del modulo Workspace puede cargar documentos para analisis."),
        ("Cuantos formatos puede generar CecilIA?", "Los 30 formatos de la Guia de Auditoria Financiera (GAF) de la CGR."),
        ("Que es el RAG?", "Retrieval-Augmented Generation: CecilIA busca en su base de documentos antes de responder, citando fuentes."),
    ]

    for pregunta, respuesta in faqs:
        doc.h3(f"P: {pregunta}")
        doc.parrafo(f"R: {respuesta}")

    # 12. Glosario
    doc.h1("Glosario de Terminos de Control Fiscal", "12.")
    terminos = [
        ("Auditoria", "Proceso sistematico de evaluacion de la gestion fiscal de una entidad"),
        ("CGR", "Contraloria General de la Republica de Colombia"),
        ("Circular 023", "Directriz del Contralor General sobre uso responsable de IA"),
        ("Condicion", "Situacion actual encontrada por el auditor en la entidad"),
        ("Control fiscal", "Funcion publica de vigilancia de la gestion fiscal del Estado"),
        ("Criterio", "Norma, ley o regla contra la cual se evalua la condicion"),
        ("DES", "Direccion de Estudios Sectoriales — control fiscal macro"),
        ("Dictamen", "Opinion del auditor sobre los estados financieros"),
        ("DVF", "Direccion de Vigilancia Fiscal — control fiscal micro"),
        ("Efecto", "Consecuencia o impacto de la irregularidad encontrada"),
        ("GAF", "Guia de Auditoria Financiera de la CGR"),
        ("Hallazgo", "Diferencia significativa entre lo esperado y lo encontrado"),
        ("ISSAI", "International Standards of Supreme Audit Institutions"),
        ("Materialidad", "Umbral de importancia relativa para errores significativos"),
        ("NIA", "Normas Internacionales de Auditoria"),
        ("RAG", "Retrieval-Augmented Generation — busqueda semantica en documentos"),
        ("Responsabilidad fiscal", "Proceso para determinar dano patrimonial"),
        ("SECOP", "Sistema Electronico de Contratacion Publica"),
    ]
    doc.tabla(["Termino", "Definicion"], [[t, d] for t, d in terminos])

    # 13. Soporte
    doc.h1("Contacto de soporte tecnico", "13.")
    doc.parrafo("Para asistencia tecnica con CecilIA v2, contacte al equipo de CD-TIC-CGR:")
    doc.tabla(
        ["Canal", "Detalle"],
        [
            ["Lider Tecnico", "Ing. Gustavo Adolfo Castillo Romero"],
            ["Correo", "soporte.cecilia@contraloria.gov.co"],
            ["Horario", "Lunes a viernes, 8:00 a 17:00"],
        ],
    )

    doc.agregar_nota_ia()
    doc.agregar_contraportada()
    return doc.guardar()


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENTO 2: Manual Tecnico
# ══════════════════════════════════════════════════════════════════════════════

def generar_manual_tecnico() -> Path:
    doc = DocCGR("Manual Tecnico de CecilIA v2",
                 "Arquitectura, deployment, API y mantenimiento",
                 "02_Manual_Tecnico_CecilIA_v2")
    doc.agregar_portada()
    doc.agregar_encabezado_pagina()

    # 1. Arquitectura
    doc.h1("Arquitectura del sistema", "1.")
    doc.h2("Componentes principales", "1.1")
    doc.parrafo(
        "CecilIA v2 utiliza una arquitectura de microservicios contenerizados con Docker Compose. "
        "Los 5 servicios principales son: PostgreSQL 16 (base de datos relacional con pgvector), "
        "Redis 7 (cache y sesiones), Backend Python (FastAPI con LangGraph), "
        "Frontend Next.js 15 (React 19 + TypeScript), y Scheduler (tareas programadas)."
    )

    doc.h2("Flujo de datos", "1.2")
    doc.parrafo(
        "El flujo tipico es: Usuario → Frontend (Next.js) → API REST (FastAPI) → Supervisor LangGraph "
        "→ Agente especializado → Tools (RAG, normativa, formatos) → Respuesta SSE streaming → Frontend. "
        "Las consultas al LLM incluyen contexto del RAG (Retrieval-Augmented Generation) con documentos "
        "normativos del corpus."
    )

    doc.h2("Diagrama de base de datos", "1.3")
    doc.parrafo("La base de datos PostgreSQL contiene 16+ tablas organizadas en modulos:")
    doc.tabla(
        ["Modulo", "Tablas", "Descripcion"],
        [
            ["Auth", "usuarios", "Autenticacion JWT y RBAC con 8 roles"],
            ["Chat", "conversaciones, mensajes", "Historial persistente de chat"],
            ["RAG", "colecciones, documentos, chunks", "Pipeline de embeddings"],
            ["Auditorias", "hallazgos, evidencias", "Gestion de hallazgos y workflow"],
            ["Capacitacion", "rutas, lecciones, progreso, quizzes", "Sistema de aprendizaje"],
            ["Capacitacion 2.0", "perfiles_aprendizaje, gamificacion, repasos, glosario", "Aprendizaje adaptativo"],
            ["Formatos", "(generados on-demand)", "30 formatos GAF en DOCX"],
            ["Observatorio", "alertas_tic", "Monitoreo sector TIC"],
            ["Analitica", "log_trazabilidad, feedback", "Metricas y retroalimentacion"],
        ],
    )

    # 2. Stack
    doc.h1("Stack tecnologico", "2.")
    doc.tabla(
        ["Componente", "Tecnologia", "Version"],
        [
            ["Backend", "Python + FastAPI", "3.12 + 0.115"],
            ["Frontend", "Next.js + React + TypeScript", "15 + 19 + 5.x"],
            ["UI", "Tailwind CSS + shadcn/ui", "3.x + latest"],
            ["BD", "PostgreSQL + pgvector", "16 + 0.7"],
            ["Cache", "Redis", "7.x"],
            ["LLM Orchestration", "LangGraph + LangChain", "0.4.7 + 0.3.25"],
            ["LLM Provider", "Multi-proveedor (Claude/Gemini/Ollama/vLLM)", "Configurable"],
            ["Embeddings", "Ollama + nomic-embed-text", "768 dims"],
            ["ORM", "SQLAlchemy (async) + asyncpg", "2.0 + 0.30"],
            ["Documentos", "python-docx + openpyxl", "1.1 + 3.1"],
            ["Audio", "edge-tts + pydub", "6.1 + 0.25"],
            ["Contenedores", "Docker + Docker Compose", "latest"],
        ],
    )

    # 3. Estructura
    doc.h1("Estructura del proyecto", "3.")
    doc.codigo(
        "cecilia-v2/\n"
        "  backend/           Python 3.12 + FastAPI\n"
        "    app/\n"
        "      agents/        LangGraph: supervisor, 5 agentes fase, transversales\n"
        "      rag/           Pipeline RAG: ingesta, chunking, embeddings\n"
        "      tools/         Tools LangChain: normativa, materialidad, Benford\n"
        "      formatos/      Generador de Formatos CGR 1-30 en DOCX\n"
        "      models/        SQLAlchemy: usuario, conversacion, hallazgo\n"
        "      api/           Endpoints REST + SSE streaming\n"
        "      auth/          JWT + RBAC 8 roles\n"
        "      services/      Logica de negocio\n"
        "  frontend/          Next.js 15 + React 19\n"
        "    app/(dashboard)/ Chat, formatos, hallazgos, capacitacion\n"
        "    lib/             api.ts, sse.ts, auth.ts, types.ts\n"
        "  corpus/            7 colecciones de documentos\n"
        "  scripts/           seed, ingest, deploy\n"
        "  docs/              Documentacion tecnica\n"
    )

    # 4. Variables de entorno
    doc.h1("Variables de entorno", "4.")
    doc.tabla(
        ["Variable", "Descripcion", "Obligatoria"],
        [
            ["LLM_BASE_URL", "URL base del proveedor LLM", "Si"],
            ["LLM_MODEL", "Nombre del modelo LLM", "Si"],
            ["LLM_API_KEY", "API key del proveedor", "Si"],
            ["DATABASE_URL", "URL de conexion PostgreSQL", "Si"],
            ["REDIS_URL", "URL de conexion Redis", "Si"],
            ["JWT_SECRET_KEY", "Clave secreta para tokens JWT", "Si"],
            ["EMBEDDINGS_MODEL", "Modelo de embeddings", "No (default: nomic-embed-text)"],
            ["LLM_TEMPERATURA", "Temperatura del modelo (0.0-1.0)", "No (default: 0.2)"],
            ["LLM_MAX_TOKENS", "Tokens maximos por respuesta", "No (default: 4096)"],
        ],
    )

    # 5. Multi-proveedor
    doc.h1("Sistema multi-proveedor LLM", "5.")
    doc.parrafo(
        "CecilIA detecta automaticamente el proveedor LLM por la URL base. "
        "Solo se cambian 3 variables en .env para cambiar de proveedor. "
        "Proveedores soportados: Anthropic (Claude), Google (Gemini), Groq, Ollama (local), vLLM."
    )

    # 6. Pipeline RAG
    doc.h1("Pipeline RAG", "6.")
    doc.h2("Ingesta de documentos", "6.1")
    doc.parrafo(
        "Formatos soportados: PDF, DOCX, TXT, CSV, XLSX. El script ingest_corpus.py procesa documentos "
        "por coleccion, aplica chunking semantico y genera embeddings con Ollama."
    )

    doc.h2("Corpus actual", "6.2")
    doc.tabla(
        ["Coleccion", "Documentos", "Chunks", "Estado"],
        [
            ["normativo", "~50", "1,824", "Poblada"],
            ["institucional", "~100", "7,840", "Poblada"],
            ["auditoria", "~20", "352", "Poblada"],
            ["academico", "-", "0", "Pendiente"],
            ["tecnico_tic", "-", "0", "Pendiente"],
            ["estadistico", "-", "0", "Pendiente"],
            ["jurisprudencial", "-", "0", "Pendiente"],
        ],
    )

    # 7. Sistema de agentes
    doc.h1("Sistema de agentes LangGraph", "7.")
    doc.parrafo(
        "CecilIA utiliza un grafo dirigido con LangGraph. El Supervisor analiza la consulta y la "
        "dirige al agente especializado apropiado. Cada agente tiene su propio system prompt y tools."
    )
    doc.tabla(
        ["Agente", "Funcion"],
        [
            ["Supervisor", "Clasifica la consulta y selecciona el agente"],
            ["Fase 0 (Pre-planeacion)", "Planeacion inicial de auditoria"],
            ["Fase 1 (Planeacion)", "Estrategia, materialidad, riesgos"],
            ["Fase 2 (Ejecucion)", "Pruebas, evidencia, hallazgos"],
            ["Fase 3 (Informe)", "Dictamen, informe de auditoria"],
            ["Fase 4 (Seguimiento)", "Planes de mejoramiento"],
            ["Analista financiero", "Analisis de estados financieros"],
            ["Analista normativo", "Busqueda y analisis de normatividad"],
            ["Generador formatos", "Creacion de formatos CGR 1-30"],
            ["Detector fraude", "Ley de Benford y analisis estadistico"],
            ["Tutor", "Capacitacion y explicaciones didacticas"],
        ],
    )

    # 8. Autenticacion
    doc.h1("Autenticacion y RBAC", "8.")
    doc.tabla(
        ["Rol", "Acceso", "Direccion"],
        [
            ["admin_tic", "Todo el sistema", "DES + DVF"],
            ["lider_tecnico", "Todo + config IA", "DES + DVF"],
            ["coordinador", "Chat, auditorias, reportes", "DES + DVF"],
            ["director_des", "Chat, auditorias, aprobar, analitica", "Solo DES"],
            ["auditor_des", "Chat, workspace, auditorias", "Solo DES"],
            ["director_dvf", "Chat, auditorias, hallazgos, analitica", "Solo DVF"],
            ["auditor_dvf", "Chat, workspace, hallazgos, formatos", "Solo DVF"],
            ["observatorio", "Solo lectura", "DES + DVF"],
            ["aprendiz", "Solo capacitacion", "DES o DVF"],
        ],
    )

    # 9-13 abreviados por espacio
    doc.h1("Guia de despliegue", "9.")
    doc.h2("Desarrollo local", "9.1")
    doc.codigo("git clone [repo]\ncd cecilia-v2\ncp .env.example .env\n# Configurar variables\ndocker compose up --build -d")

    doc.h2("Produccion con Nginx y SSL", "9.2")
    doc.parrafo(
        "En produccion se utiliza Nginx como reverse proxy con certificados SSL. "
        "La configuracion se encuentra en infra/nginx/. El despliegue se realiza con scripts/deploy.sh."
    )

    doc.h1("Troubleshooting", "10.")
    doc.tabla(
        ["Problema", "Solucion"],
        [
            ["Backend no inicia", "Verificar DATABASE_URL y que PostgreSQL este corriendo"],
            ["LLM no responde", "Verificar LLM_BASE_URL, LLM_API_KEY y creditos disponibles"],
            ["Frontend no carga", "Verificar NEXT_PUBLIC_API_URL apunta al backend"],
            ["Audio no se genera", "Verificar edge-tts y ffmpeg instalados en container"],
            ["Formatos sin logo", "Verificar logo-cgr.png en backend/app/static/"],
        ],
    )

    doc.agregar_nota_ia()
    doc.agregar_contraportada()
    return doc.guardar()


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENTOS 3-15: Politicas institucionales
# ══════════════════════════════════════════════════════════════════════════════

POLITICAS = [
    ("03", "Gobernanza de IA para Control Fiscal", "gobernanza_ia",
     "Establece el marco de gobierno para el uso de inteligencia artificial en la funcion de control fiscal.",
     [
         ("Objeto y alcance", "Esta politica establece los principios, roles y mecanismos de gobernanza para el uso responsable de inteligencia artificial en CecilIA v2, sistema de IA para control fiscal de la CGR."),
         ("Marco normativo", "Circular 023 de la CGR (2025IE0146473), CONPES 4144/2025 (Politica Nacional de IA), CONPES 4145/2025 (IA para el Estado), Ley 1581/2012 (Proteccion datos personales), NIST AI RMF 1.0, UNESCO Etica IA 2021, OCDE Principios IA 2019, ISO/IEC 42001."),
         ("Principios de gobernanza", "Supervision humana obligatoria: Ninguna decision fiscal se toma automaticamente. Transparencia: Toda respuesta de CecilIA cita sus fuentes. Trazabilidad: Cada consulta queda registrada en log_trazabilidad. Proporcionalidad: El uso de IA se limita a asistencia, nunca a reemplazo."),
         ("Estructura de gobernanza", "Comite de Gobernanza IA: Conformado por Contralor Delegado TIC (presidente), directores DES y DVF, coordinadora juridica, y lider tecnico. Se reune trimestralmente para evaluar desempeno, riesgos y actualizaciones del sistema."),
         ("Ciclo de vida del modelo", "Evaluacion pre-despliegue: Benchmark de 50 preguntas con metricas de precision, recall y calidad. Monitoreo continuo: Dashboard de analitica con feedback de usuarios. Actualizacion: Protocolo de actualizacion del corpus normativo cada vez que se emita nueva normatividad."),
     ]),
    ("04", "Supervision Humana Obligatoria", "supervision_humana",
     "Garantiza que todo producto generado por CecilIA sea revisado y validado por un auditor.",
     [
         ("Objeto y alcance", "Definir los mecanismos de supervision humana para todo contenido generado por CecilIA v2, conforme a la Circular 023 de la CGR y los principios de la UNESCO sobre IA."),
         ("Marco normativo", "Circular 023 CGR — Principio 4: Supervision humana. CONPES 4144/2025 — Eje 3: Confianza. Sentencia T-323/2024 Corte Constitucional — Criterios uso responsable IA."),
         ("Niveles de supervision", "Nivel 1 (Chat): El auditor evalua cada respuesta antes de usarla. Nivel 2 (Formatos): Checkbox obligatorio de validacion antes de descarga. Nivel 3 (Hallazgos): Workflow de aprobacion con firma del director."),
         ("Usos limitados por Circular 023", "CecilIA NO puede generar como producto final: dictamenes, opiniones sobre estados contables, hallazgos definitivos, diagnosticos sectoriales completos ni informes finales. CecilIA SOLO genera borradores y sugerencias."),
         ("Controles implementados", "Disclaimer automatico en cada respuesta del chat. Modal de aceptacion al primer login. Checkbox de validacion en formatos. Nota de IA en pie de pagina de todo DOCX."),
     ]),
    ("05", "Transparencia y Explicabilidad Algoritmica", "transparencia",
     "Asegura que los usuarios comprendan como CecilIA genera sus respuestas.",
     [
         ("Objeto y alcance", "Establecer los mecanismos de transparencia y explicabilidad para que los auditores comprendan como CecilIA genera respuestas, de donde obtiene la informacion y cuales son sus limitaciones."),
         ("Marco normativo", "CONPES 4144/2025 — Eje de transparencia. UNESCO Etica IA 2021 — Principio de explicabilidad. OCDE Principios IA 2019 — Transparencia y rendicion de cuentas. ISO/IEC 42001 — Clausula 7.4 Comunicacion."),
         ("Mecanismos de transparencia", "Citacion de fuentes: Cada respuesta incluye las fuentes RAG consultadas. Modelo visible: El header muestra el modelo LLM activo. Trazabilidad: Log completo de cada interaccion almacenado. Pagina /acerca: Informacion publica del sistema."),
         ("Explicabilidad tecnica", "El sistema RAG permite explicar por que se genero una respuesta: se muestran los chunks de texto recuperados, la coleccion de origen y el score de relevancia."),
         ("Limitaciones conocidas", "Alucinaciones: El LLM puede generar informacion plausible pero incorrecta. Sesgo: Los modelos reflejan sesgos de sus datos de entrenamiento. Actualizacion: El corpus puede no incluir normatividad emitida recientemente."),
     ]),
    ("06", "Calidad y Gobernanza de Datos", "calidad_datos",
     "Define estandares para la gestion de datos en CecilIA.",
     [
         ("Objeto y alcance", "Establecer los estandares de calidad, integridad y gobernanza de los datos utilizados por CecilIA v2, incluyendo el corpus normativo, datos de auditoria y registros del sistema."),
         ("Marco normativo", "Ley 1581/2012 (Proteccion de datos). Decreto 1377/2013 (Reglamentario). CONPES 4144/2025. ISO/IEC 42001 — Clausula 6.1.2 Datos de IA."),
         ("Clasificacion de datos", "Datos publicos: Normatividad, jurisprudencia. Datos internos: Parametros de auditoria, formatos. Datos sensibles: Hallazgos en proceso, datos de entidades. Datos personales: Informacion de usuarios del sistema."),
         ("Estandares de calidad del corpus", "Autenticidad: Solo documentos de fuentes oficiales. Vigencia: Normatividad actualizada. Integridad: Verificacion de completitud post-ingesta. Trazabilidad: Metadatos de cada documento (fuente, fecha, responsable)."),
         ("Anonimizacion", "Conforme a la Ley 1581/2012, todos los datos personales en ejemplos y entrenamiento deben ser anonimizados. CecilIA detecta automaticamente datos personales en consultas y genera alertas de privacidad."),
     ]),
    ("07", "Ciberseguridad en Sistemas de IA", "ciberseguridad",
     "Protege CecilIA contra amenazas de ciberseguridad.",
     [
         ("Objeto y alcance", "Definir los controles de ciberseguridad para proteger CecilIA v2 contra amenazas, incluyendo inyeccion de prompts, exfiltracion de datos y ataques adversariales."),
         ("Marco normativo", "NIST AI RMF 1.0 — Categoria SECURE. ISO/IEC 42001 — Anexo A.6 Seguridad. CONPES 4144/2025 — Eje de seguridad."),
         ("Controles implementados", "Autenticacion: JWT con bcrypt hash de passwords. Cifrado: AES-256 en reposo, TLS 1.3 en transito. RBAC: 8 roles con permisos granulares. API keys: Solo en .env, nunca en codigo. Logs: Trazabilidad completa de cada consulta."),
         ("Amenazas especificas de IA", "Prompt injection: Validacion y sanitizacion de inputs. Data poisoning: Control de calidad del corpus. Model extraction: Rate limiting y monitoreo. Privacy leakage: Anonimizacion y filtros de salida."),
         ("Plan de respuesta a incidentes", "Deteccion: Monitoreo automatico de anomalias. Contencion: Desactivacion inmediata del servicio afectado. Erradicacion: Identificacion y correccion de la vulnerabilidad. Recuperacion: Restauracion desde backups verificados."),
     ]),
    ("08", "Etica, Integridad Cientifica y Principios FAIR", "etica_fair",
     "Marco etico para el desarrollo y uso de CecilIA.",
     [
         ("Objeto y alcance", "Establecer el marco etico que rige el desarrollo, despliegue y uso de CecilIA v2, basado en principios internacionales de etica en IA."),
         ("Marco normativo", "UNESCO Recomendacion sobre Etica de la IA (2021). OCDE Principios IA (2019). CONPES 4144/2025 — Eje etico. Sentencia T-323/2024."),
         ("Principios eticos", "Beneficencia: CecilIA debe contribuir al bien publico. No maleficencia: Minimizar riesgos de dano. Autonomia: El auditor mantiene la decision final. Justicia: Acceso equitativo sin discriminacion. Explicabilidad: Toda respuesta es trazable."),
         ("Principios FAIR", "Findable: Los datos del corpus son catalogados con metadatos. Accessible: Acceso controlado por roles. Interoperable: Formatos estandar (DOCX, JSON, CSV). Reusable: Documentacion completa del corpus."),
         ("Comite de etica", "Se recomienda conformar un comite de etica IA que evalue periodicamente el impacto del sistema en los procesos de control fiscal."),
     ]),
    ("09", "Propiedad Intelectual del Software", "propiedad_intelectual",
     "Protege los derechos de autor sobre CecilIA v2.",
     [
         ("Objeto y alcance", "Establecer el regimen de propiedad intelectual aplicable a CecilIA v2 como obra de software desarrollada por la CGR."),
         ("Marco normativo", "Ley 23 de 1982 (Derechos de autor). Decision Andina 351 de 1993. Ley 1915 de 2018. Decreto 1360 de 1989."),
         ("Titularidad", "Los derechos patrimoniales de CecilIA v2 pertenecen a la Contraloria General de la Republica de Colombia, como obra creada por funcionarios en ejercicio de sus funciones."),
         ("Registro DNDA", "El software sera registrado ante la Direccion Nacional de Derechos de Autor (DNDA) conforme al procedimiento establecido."),
         ("Dependencias open source", "CecilIA v2 utiliza librerias de codigo abierto con licencias compatibles (MIT, Apache 2.0, BSD). Se mantiene un inventario completo de dependencias y sus licencias."),
     ]),
    ("10", "Sostenibilidad Ambiental en IA", "sostenibilidad",
     "Minimiza el impacto ambiental del uso de IA en la CGR.",
     [
         ("Objeto y alcance", "Definir estrategias para minimizar el impacto ambiental del despliegue y operacion de CecilIA v2."),
         ("Marco normativo", "CONPES 4144/2025 — Eje de sostenibilidad. ODS 13 (Accion por el clima). UNESCO Etica IA 2021 — Principio de sostenibilidad."),
         ("Estrategias implementadas", "Modelos eficientes: Uso preferente de modelos compactos (Gemini Flash, Ollama local). Cache inteligente: Redis para evitar consultas repetidas al LLM. Embeddings locales: Ollama en servidor propio reduce trafico externo. Monitoreo de consumo: Metricas de tokens consumidos por usuario."),
         ("Medicion", "Se registra el consumo de tokens, tiempo de CPU y ancho de banda por cada interaccion para evaluar la huella computacional del sistema."),
     ]),
    ("11", "Uso Responsable de CecilIA", "uso_responsable",
     "Guia para el uso etico y responsable del sistema por parte de los usuarios.",
     [
         ("Objeto y alcance", "Establecer las normas de uso responsable de CecilIA v2 para todos los funcionarios de la CGR autorizados."),
         ("Marco normativo", "Circular 023 CGR (todos los 10 principios). CONPES 4144/2025. Codigo Unico Disciplinario."),
         ("Reglas de uso", "Validacion obligatoria: Todo contenido generado debe ser verificado antes de uso oficial. Confidencialidad: No ingresar datos clasificados o de seguridad nacional. Anonimizacion: No incluir datos personales reales en las consultas. Feedback: Reportar errores o alucinaciones mediante el sistema de retroalimentacion."),
         ("Usos prohibidos", "Generar documentos oficiales sin revision humana. Tomar decisiones fiscales basadas unicamente en respuestas de IA. Compartir credenciales de acceso. Usar CecilIA para fines personales o ajenos al control fiscal."),
         ("Sanciones", "El uso indebido de CecilIA puede acarrear sanciones disciplinarias conforme al Codigo Unico Disciplinario (Ley 1952/2019)."),
     ]),
    ("12", "Gestion de Incidentes de IA", "gestion_incidentes",
     "Protocolo para responder a incidentes relacionados con el sistema de IA.",
     [
         ("Objeto y alcance", "Definir el protocolo de respuesta ante incidentes relacionados con CecilIA v2, incluyendo alucinaciones criticas, sesgos detectados y fallos de seguridad."),
         ("Marco normativo", "NIST AI RMF 1.0 — RESPOND. ISO/IEC 42001 — Clausula 10 Mejora."),
         ("Clasificacion de incidentes", "Critico: Respuesta que podria causar dano fiscal si se usa sin validar. Alto: Alucinacion en tema normativo que contradice la ley vigente. Medio: Error factual en datos no criticos. Bajo: Error de formato o presentacion."),
         ("Protocolo de respuesta", "1. Deteccion: Via feedback de usuario o monitoreo automatico. 2. Escalamiento: Segun nivel de criticidad al equipo tecnico. 3. Contencion: Desactivacion del modulo afectado si es necesario. 4. Correccion: Actualizacion del corpus o ajuste del sistema. 5. Documentacion: Registro del incidente y leccion aprendida."),
     ]),
    ("13", "Evaluacion de Impacto Algoritmico", "evaluacion_impacto",
     "Metodologia para evaluar el impacto de CecilIA en los procesos de control fiscal.",
     [
         ("Objeto y alcance", "Establecer la metodologia para evaluar periodicamente el impacto de CecilIA v2 en la calidad, eficiencia y equidad del control fiscal."),
         ("Marco normativo", "NIST AI RMF 1.0 — MAP/MEASURE. ISO/IEC 42001 — Clausula 9 Evaluacion del desempeno."),
         ("Indicadores de impacto", "Eficiencia: Tiempo de generacion de formatos vs manual. Calidad: Precision de respuestas medida con benchmark. Cobertura: Porcentaje de consultas resueltas satisfactoriamente. Adopcion: Numero de usuarios activos y frecuencia de uso."),
         ("Metodologia", "Evaluacion trimestral con benchmark de 50 preguntas. Encuesta de satisfaccion a usuarios. Analisis de feedback acumulado. Comparativo con procesos sin IA."),
     ]),
    ("14", "Interoperabilidad con Sistemas CGR", "interoperabilidad",
     "Garantiza la integracion de CecilIA con los sistemas existentes de la CGR.",
     [
         ("Objeto y alcance", "Definir los estandares y protocolos de integracion de CecilIA v2 con los sistemas de informacion existentes en la CGR."),
         ("Marco normativo", "CONPES 4145/2025 — Interoperabilidad estatal. Marco de Interoperabilidad del Estado colombiano."),
         ("Sistemas integrados", "SIRECI: Sistema de Rendicion Electronica de Cuenta e Informes (stub). SIGECI: Sistema Integrado de Gestion (stub). APA: Aplicativo de Plan de Accion (stub). DIARI: Datos de auditoria (stub). SECOP: Integracion activa con API publica."),
         ("APIs externas activas", "SECOP (dane.gov.co): Contratos publicos. Congreso (congresovisible.uniandes.edu.co): Proyectos de ley. DANE (dane.gov.co): Indicadores economicos."),
         ("Estandares", "Formato JSON para intercambio de datos. API REST con documentacion OpenAPI. Autenticacion OAuth2/JWT."),
     ]),
    ("15", "Escalabilidad a Otras Delegadas", "escalabilidad",
     "Plan para replicar CecilIA en otras contralorias delegadas.",
     [
         ("Objeto y alcance", "Definir la estrategia y requisitos para escalar CecilIA v2 a otras contralorias delegadas de la CGR y eventualmente a contralorias territoriales."),
         ("Marco normativo", "CONPES 4145/2025 — IA para el Estado. Decreto 2037 de 2019 — Estructura CGR."),
         ("Modelo de escalabilidad", "Fase 1: Piloto en CD-TIC con DES y DVF (actual). Fase 2: Extension a delegadas sectoriales (agricultura, infraestructura, social). Fase 3: Contralorias territoriales via SaaS."),
         ("Requisitos por delegada", "Servidor con 16GB RAM minimo. Acceso a red CGR. Corpus normativo especifico de la delegada. Capacitacion de usuarios (4-8 horas)."),
         ("Personalizacion", "Cada delegada configura: su corpus normativo, sus formatos especificos, y sus agentes especializados. El core del sistema permanece igual."),
     ]),
]


def generar_politica(num: str, titulo: str, slug: str, desc: str, secciones: list) -> Path:
    """Genera un documento de politica institucional."""
    doc = DocCGR(f"Politica de {titulo}",
                 desc,
                 f"{num}_Politica_{slug}")
    doc.agregar_portada()
    doc.agregar_encabezado_pagina()

    for i, (sec_titulo, sec_contenido) in enumerate(secciones, 1):
        doc.h1(sec_titulo, f"{i}.")
        # Dividir contenido largo en parrafos
        partes = sec_contenido.split(". ")
        texto_actual = ""
        for parte in partes:
            texto_actual += parte + ". "
            if len(texto_actual) > 300:
                doc.parrafo(texto_actual.strip())
                texto_actual = ""
        if texto_actual.strip():
            doc.parrafo(texto_actual.strip())

    # Roles y responsabilidades
    doc.h1("Roles y responsabilidades", f"{len(secciones) + 1}.")
    doc.tabla(
        ["Rol", "R", "A", "C", "I"],
        [
            ["Contralor Delegado TIC", "", "A", "", ""],
            ["Director DES/DVF", "", "", "C", "I"],
            ["Lider Tecnico", "R", "", "", ""],
            ["Coordinadora Juridica", "", "", "C", ""],
            ["Auditores", "", "", "", "I"],
        ],
    )

    # Indicadores
    doc.h1("Indicadores de cumplimiento", f"{len(secciones) + 2}.")
    doc.parrafo("Se mediran trimestralmente los siguientes indicadores:")
    doc.lista([
        "Porcentaje de cumplimiento de los controles establecidos",
        "Numero de incidentes reportados relacionados con esta politica",
        "Resultado de evaluaciones periodicas del equipo tecnico",
    ])

    # Revision
    doc.h1("Periodicidad de revision", f"{len(secciones) + 3}.")
    doc.parrafo("Esta politica sera revisada semestralmente o cuando se presenten cambios normativos significativos.")

    # Control de versiones
    doc.h1("Control de versiones", f"{len(secciones) + 4}.")
    doc.tabla(
        ["Version", "Fecha", "Descripcion", "Elaborado por"],
        [["1.0", "Abril 2026", "Version inicial", "Equipo Tecnico CecilIA"]],
    )

    doc.agregar_nota_ia()
    doc.agregar_contraportada()
    return doc.guardar()


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENTO 16: Registro DNDA
# ══════════════════════════════════════════════════════════════════════════════

def generar_registro_dnda() -> Path:
    doc = DocCGR("Registro de Software ante DNDA",
                 "Direccion Nacional de Derechos de Autor",
                 "16_Registro_Software_DNDA_CecilIA")
    doc.agregar_portada()
    doc.agregar_encabezado_pagina()

    doc.h1("Datos del programa de computador", "1.")
    doc.tabla(
        ["Campo", "Valor"],
        [
            ["Nombre del programa", "CecilIA v2 — Sistema de IA para Control Fiscal"],
            ["Descripcion", "Sistema de inteligencia artificial que asiste a auditores de la CGR"],
            ["Lenguajes", "Python 3.12, TypeScript 5.x, JavaScript, SQL, HTML, CSS"],
            ["Fecha de creacion", "Enero 2026"],
            ["Fecha primera publicacion", "Abril 2026"],
            ["Tipo de obra", "Programa de computador (software)"],
            ["Pais de origen", "Colombia"],
        ],
    )

    doc.h1("Titular de derechos patrimoniales", "2.")
    doc.parrafo("Contraloria General de la Republica de Colombia")
    doc.parrafo("NIT: 899.999.067-3")
    doc.parrafo("Direccion: Carrera 69 No. 44-35, Bogota D.C., Colombia")

    doc.h1("Autores", "3.")
    doc.tabla(
        ["Nombre", "Cargo", "Contribucion"],
        [[n, c, r] for n, c, r in EQUIPO],
    )

    doc.h1("Modulos del sistema", "4.")
    doc.tabla(
        ["Modulo", "Descripcion funcional"],
        [
            ["Chat IA", "Interfaz conversacional con streaming SSE, multi-agente"],
            ["Formatos CGR", "Generacion de 30 formatos GAF en DOCX institucional"],
            ["Hallazgos", "CRUD con 4 elementos, connotaciones y workflow"],
            ["Capacitacion", "Rutas adaptativas, gamificacion, tutor IA"],
            ["Observatorio TIC", "Monitoreo y alertas del sector TIC"],
            ["Analitica", "Dashboard ejecutivo con metricas y feedback"],
            ["Workspace", "Carga y analisis de documentos"],
            ["RAG", "Pipeline de ingesta, embeddings y busqueda semantica"],
        ],
    )

    doc.h1("Dependencias de codigo abierto", "5.")
    doc.tabla(
        ["Libreria", "Licencia", "Uso"],
        [
            ["FastAPI", "MIT", "Framework web backend"],
            ["Next.js", "MIT", "Framework frontend"],
            ["React", "MIT", "Biblioteca de UI"],
            ["LangChain", "MIT", "Orquestacion LLM"],
            ["LangGraph", "MIT", "Grafo de agentes"],
            ["SQLAlchemy", "MIT", "ORM de base de datos"],
            ["python-docx", "MIT", "Generacion de DOCX"],
            ["edge-tts", "MIT", "Sintesis de voz"],
            ["Tailwind CSS", "MIT", "Framework CSS"],
            ["PostgreSQL", "PostgreSQL License", "Base de datos"],
            ["Redis", "BSD-3", "Cache y sesiones"],
        ],
    )

    doc.agregar_nota_ia()
    doc.agregar_contraportada()
    return doc.guardar()


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENTO 17: Guia de implementacion
# ══════════════════════════════════════════════════════════════════════════════

def generar_guia_implementacion() -> Path:
    doc = DocCGR("Guia de Implementacion Institucional",
                 "Para replicar CecilIA en otras delegadas y contralorias",
                 "17_Guia_Implementacion_Institucional_CecilIA")
    doc.agregar_portada()
    doc.agregar_encabezado_pagina()

    doc.h1("Requisitos de infraestructura", "1.")
    doc.tabla(
        ["Componente", "Minimo", "Recomendado"],
        [
            ["CPU", "4 cores", "8 cores"],
            ["RAM", "16 GB", "32 GB"],
            ["Disco", "100 GB SSD", "500 GB SSD"],
            ["Red", "100 Mbps", "1 Gbps"],
            ["SO", "Ubuntu 22.04 LTS", "Ubuntu 22.04 LTS"],
            ["Docker", "24.x", "Ultima version estable"],
        ],
    )

    doc.h1("Proceso de instalacion", "2.")
    doc.lista([
        "Provisionar servidor con requisitos minimos",
        "Instalar Docker y Docker Compose",
        "Clonar repositorio cecilia-v2",
        "Configurar variables de entorno (.env)",
        "Ejecutar docker compose up --build -d",
        "Ejecutar seed de usuarios: python scripts/seed_users.py",
        "Ingestar corpus normativo: python scripts/ingest_corpus.py",
        "Configurar Nginx con SSL para acceso externo",
        "Verificar salud de servicios en /salud",
    ], ordenada=True)

    doc.h1("Configuracion del corpus", "3.")
    doc.parrafo(
        "Cada delegada debe preparar su corpus normativo especifico. Los documentos se organizan "
        "en 7 colecciones: normativo, institucional, auditoria, academico, tecnico_tic, estadistico, "
        "jurisprudencial. El script de ingesta procesa PDF, DOCX, TXT y CSV."
    )

    doc.h1("Capacitacion de usuarios", "4.")
    doc.tabla(
        ["Sesion", "Duracion", "Contenido"],
        [
            ["Introduccion", "2 horas", "Vision general, principios Circular 023"],
            ["Chat y RAG", "2 horas", "Como hacer consultas efectivas"],
            ["Formatos y hallazgos", "2 horas", "Generacion de DOCX, workflow"],
            ["Administracion", "2 horas", "Gestion de usuarios, corpus, monitoreo"],
        ],
    )

    doc.h1("Cronograma estimado", "5.")
    doc.tabla(
        ["Semana", "Actividad"],
        [
            ["1", "Provision de infraestructura y configuracion"],
            ["2", "Instalacion y pruebas de funcionamiento"],
            ["3", "Ingesta del corpus normativo de la delegada"],
            ["4", "Capacitacion de usuarios y pruebas piloto"],
            ["5-6", "Ajustes y puesta en produccion"],
        ],
    )

    doc.h1("Soporte y mantenimiento", "6.")
    doc.parrafo(
        "El equipo de CD-TIC proporcionara soporte nivel 1 y 2 durante los primeros 3 meses. "
        "Cada delegada designara un administrador local capacitado para soporte nivel 0 (usuario). "
        "Las actualizaciones del sistema se distribuiran via Docker images versionadas."
    )

    doc.agregar_nota_ia()
    doc.agregar_contraportada()
    return doc.guardar()


# ══════════════════════════════════════════════════════════════════════════════
# XLSX: Matriz RACI, Inventario API, Inventario Corpus
# ══════════════════════════════════════════════════════════════════════════════

def _estilo_xlsx_header():
    return {
        "font": Font(name="Calibri", size=10, bold=True, color="FFFFFF"),
        "fill": PatternFill(start_color="1A5276", end_color="1A5276", fill_type="solid"),
        "alignment": Alignment(horizontal="center", vertical="center", wrap_text=True),
        "border": Border(
            left=Side(style="thin"), right=Side(style="thin"),
            top=Side(style="thin"), bottom=Side(style="thin"),
        ),
    }


def _estilo_xlsx_celda(alterna=False):
    fill = PatternFill(start_color="EBF5FB", end_color="EBF5FB", fill_type="solid") if alterna else PatternFill()
    return {
        "font": Font(name="Calibri", size=9),
        "alignment": Alignment(vertical="center", wrap_text=True),
        "border": Border(
            left=Side(style="thin"), right=Side(style="thin"),
            top=Side(style="thin"), bottom=Side(style="thin"),
        ),
        "fill": fill,
    }


def _aplicar_estilo(celda, estilo):
    for k, v in estilo.items():
        setattr(celda, k, v)


def generar_xlsx_raci() -> Path:
    if not TIENE_OPENPYXL:
        return Path("N/A")

    wb = Workbook()

    # Hoja 1: RACI
    ws = wb.active
    ws.title = "RACI"
    encabezados = ["Actividad", "Contralor Delegado", "Director DES/DVF", "Lider Tecnico", "Coordinador", "Auditores"]
    for j, enc in enumerate(encabezados, 1):
        c = ws.cell(row=1, column=j, value=enc)
        _aplicar_estilo(c, _estilo_xlsx_header())

    actividades = [
        ("Definir alcance de uso de IA", "A", "C", "R", "C", "I"),
        ("Configurar y desplegar sistema", "I", "I", "R", "I", ""),
        ("Ingestar corpus normativo", "I", "C", "R", "I", ""),
        ("Crear/gestionar usuarios", "A", "C", "R", "", ""),
        ("Supervisar respuestas de IA", "I", "A", "C", "C", "R"),
        ("Generar formatos CGR", "", "A", "", "", "R"),
        ("Aprobar hallazgos", "", "A", "", "C", "R"),
        ("Evaluar desempeno del sistema", "A", "C", "R", "C", "I"),
        ("Actualizar politicas de IA", "A", "C", "R", "R", "I"),
        ("Capacitar nuevos usuarios", "I", "C", "R", "R", "I"),
        ("Monitorear seguridad", "I", "I", "R", "C", ""),
        ("Reportar incidentes", "I", "I", "R", "C", "R"),
        ("Generar informes de analitica", "I", "A", "R", "C", "I"),
    ]

    for i, (act, *roles) in enumerate(actividades, 2):
        ws.cell(row=i, column=1, value=act)
        for j, rol in enumerate(roles, 2):
            c = ws.cell(row=i, column=j, value=rol)
            _aplicar_estilo(c, _estilo_xlsx_celda(i % 2 == 0))
        _aplicar_estilo(ws.cell(row=i, column=1), _estilo_xlsx_celda(i % 2 == 0))

    ws.column_dimensions["A"].width = 40
    for col in "BCDEF":
        ws.column_dimensions[col].width = 20

    # Hoja 2: Roles
    ws2 = wb.create_sheet("Roles")
    for j, enc in enumerate(["Rol", "Descripcion"], 1):
        _aplicar_estilo(ws2.cell(row=1, column=j, value=enc), _estilo_xlsx_header())
    roles_desc = [
        ("Contralor Delegado TIC", "Maxima autoridad del proyecto, aprueba politicas"),
        ("Director DES/DVF", "Supervision de auditoria, aprobacion de hallazgos"),
        ("Lider Tecnico", "Arquitectura, desarrollo, despliegue del sistema"),
        ("Coordinador", "Gestion de politicas, reportes, coordinacion"),
        ("Auditores", "Usuarios principales del sistema, generan formatos y hallazgos"),
    ]
    for i, (rol, desc) in enumerate(roles_desc, 2):
        ws2.cell(row=i, column=1, value=rol)
        ws2.cell(row=i, column=2, value=desc)
    ws2.column_dimensions["A"].width = 30
    ws2.column_dimensions["B"].width = 60

    # Hoja 3: Leyenda
    ws3 = wb.create_sheet("Leyenda")
    for j, enc in enumerate(["Letra", "Significado", "Descripcion"], 1):
        _aplicar_estilo(ws3.cell(row=1, column=j, value=enc), _estilo_xlsx_header())
    leyenda = [
        ("R", "Responsible", "Ejecuta la actividad"),
        ("A", "Accountable", "Aprueba y rinde cuentas"),
        ("C", "Consulted", "Aporta informacion o asesoria"),
        ("I", "Informed", "Recibe informacion del resultado"),
    ]
    for i, (l, s, d) in enumerate(leyenda, 2):
        ws3.cell(row=i, column=1, value=l)
        ws3.cell(row=i, column=2, value=s)
        ws3.cell(row=i, column=3, value=d)

    ruta = CARPETA_SALIDA / "XLSX_01_Matriz_RACI_CecilIA.xlsx"
    wb.save(str(ruta))
    return ruta


def generar_xlsx_api() -> Path:
    if not TIENE_OPENPYXL:
        return Path("N/A")

    wb = Workbook()
    ws = wb.active
    ws.title = "Endpoints"

    encabezados = ["Metodo", "Ruta", "Descripcion", "Auth", "Roles", "Request", "Response"]
    for j, enc in enumerate(encabezados, 1):
        _aplicar_estilo(ws.cell(row=1, column=j, value=enc), _estilo_xlsx_header())

    endpoints = [
        ("POST", "/auth/login", "Iniciar sesion", "No", "Todos", "username, password", "token, usuario"),
        ("GET", "/salud", "Health check", "No", "Todos", "-", "status, version"),
        ("POST", "/chat/enviar", "Enviar mensaje (SSE)", "Si", "Todos", "mensaje, conversacion_id", "SSE stream"),
        ("GET", "/conversaciones", "Listar conversaciones", "Si", "Todos", "-", "lista conversaciones"),
        ("POST", "/conversaciones", "Crear conversacion", "Si", "Todos", "titulo", "conversacion"),
        ("DELETE", "/conversaciones/{id}", "Eliminar conversacion", "Si", "Todos", "-", "ok"),
        ("GET", "/capacitacion/rutas", "Listar rutas aprendizaje", "Si", "Todos", "?direccion=", "lista rutas"),
        ("GET", "/capacitacion/lecciones/{id}", "Obtener leccion", "Si", "Todos", "-", "leccion con MD"),
        ("POST", "/capacitacion/generar-audio", "Generar podcast", "Si", "Todos", "tema, duracion", "script, audio_base64"),
        ("POST", "/capacitacion/generar-flashcards", "Generar flashcards", "Si", "Todos", "tema, num_tarjetas", "tarjetas"),
        ("POST", "/capacitacion/generar-infografia", "Generar infografia", "Si", "Todos", "tema", "mermaid, imagen_url"),
        ("POST", "/capacitacion/generar-manual", "Generar manual DOCX", "Si", "Todos", "tema, nivel", "contenido_base64"),
        ("POST", "/formatos/generar", "Generar formato CGR", "Si", "DVF", "numero, datos", "contenido_base64"),
        ("GET", "/hallazgos", "Listar hallazgos", "Si", "DVF", "?estado=", "lista hallazgos"),
        ("POST", "/hallazgos", "Crear hallazgo", "Si", "DVF", "condicion, criterio, ...", "hallazgo"),
        ("GET", "/observatorio/alertas", "Listar alertas TIC", "Si", "DES", "?nivel=", "lista alertas"),
        ("GET", "/analitica/metricas", "Dashboard metricas", "Si", "Directores", "?periodo=", "metricas"),
        ("POST", "/workspace/subir", "Subir documento", "Si", "Todos", "archivo (multipart)", "documento_id"),
    ]

    for i, (met, ruta, desc, auth, roles, req, resp) in enumerate(endpoints, 2):
        for j, val in enumerate([met, ruta, desc, auth, roles, req, resp], 1):
            c = ws.cell(row=i, column=j, value=val)
            _aplicar_estilo(c, _estilo_xlsx_celda(i % 2 == 0))

    ws.column_dimensions["A"].width = 10
    ws.column_dimensions["B"].width = 35
    ws.column_dimensions["C"].width = 30
    ws.column_dimensions["D"].width = 8
    ws.column_dimensions["E"].width = 15
    ws.column_dimensions["F"].width = 25
    ws.column_dimensions["G"].width = 25

    # Hoja 2: Codigos error
    ws2 = wb.create_sheet("Codigos Error")
    for j, enc in enumerate(["Codigo", "Descripcion", "Solucion"], 1):
        _aplicar_estilo(ws2.cell(row=1, column=j, value=enc), _estilo_xlsx_header())
    errores = [
        ("400", "Bad Request — Parametros invalidos", "Verificar formato del request body"),
        ("401", "Unauthorized — Token invalido o expirado", "Iniciar sesion nuevamente"),
        ("403", "Forbidden — Sin permisos para este recurso", "Verificar rol del usuario"),
        ("404", "Not Found — Recurso no encontrado", "Verificar ID o ruta del endpoint"),
        ("422", "Validation Error — Datos no pasan validacion", "Revisar tipos y campos requeridos"),
        ("500", "Internal Server Error", "Revisar logs del backend: docker logs cecilia-backend"),
        ("503", "Service Unavailable — LLM no disponible", "Verificar LLM_BASE_URL y creditos"),
    ]
    for i, (cod, desc, sol) in enumerate(errores, 2):
        ws2.cell(row=i, column=1, value=cod)
        ws2.cell(row=i, column=2, value=desc)
        ws2.cell(row=i, column=3, value=sol)
    ws2.column_dimensions["A"].width = 10
    ws2.column_dimensions["B"].width = 40
    ws2.column_dimensions["C"].width = 50

    ruta = CARPETA_SALIDA / "XLSX_02_Inventario_API_CecilIA.xlsx"
    wb.save(str(ruta))
    return ruta


def generar_xlsx_corpus() -> Path:
    if not TIENE_OPENPYXL:
        return Path("N/A")

    wb = Workbook()
    ws = wb.active
    ws.title = "Documentos"

    encabezados = ["Nombre", "Coleccion", "Tipo", "Chunks", "Estado"]
    for j, enc in enumerate(encabezados, 1):
        _aplicar_estilo(ws.cell(row=1, column=j, value=enc), _estilo_xlsx_header())

    docs = [
        ("Constitucion Politica de Colombia", "normativo", "PDF", "~200", "Ingestado"),
        ("Ley 42 de 1993", "normativo", "PDF", "~80", "Ingestado"),
        ("Ley 610 de 2000", "normativo", "PDF", "~120", "Ingestado"),
        ("Decreto 403 de 2020", "normativo", "PDF", "~100", "Ingestado"),
        ("CONPES 4144/2025", "normativo", "PDF", "~90", "Ingestado"),
        ("CONPES 4145/2025", "normativo", "PDF", "~85", "Ingestado"),
        ("Circular 023 CGR", "normativo", "PDF", "~30", "Ingestado"),
        ("Guia Auditoria Financiera (GAF)", "institucional", "PDF", "~2000", "Ingestado"),
        ("Manual SIGECI", "institucional", "PDF", "~500", "Ingestado"),
        ("Instructivos de formatos", "institucional", "PDF", "~800", "Ingestado"),
        ("Hallazgos de referencia", "auditoria", "PDF", "~352", "Ingestado"),
    ]

    for i, (nom, col, tipo, chunks, est) in enumerate(docs, 2):
        for j, val in enumerate([nom, col, tipo, chunks, est], 1):
            c = ws.cell(row=i, column=j, value=val)
            _aplicar_estilo(c, _estilo_xlsx_celda(i % 2 == 0))

    for col_letter, width in [("A", 40), ("B", 15), ("C", 10), ("D", 10), ("E", 12)]:
        ws.column_dimensions[col_letter].width = width

    # Hoja 2: Estadisticas
    ws2 = wb.create_sheet("Estadisticas")
    for j, enc in enumerate(["Coleccion", "Total docs", "Total chunks", "Estado"], 1):
        _aplicar_estilo(ws2.cell(row=1, column=j, value=enc), _estilo_xlsx_header())
    stats = [
        ("normativo", "~50", "1,824", "Poblada"),
        ("institucional", "~100", "7,840", "Poblada"),
        ("auditoria", "~20", "352", "Poblada"),
        ("academico", "0", "0", "Pendiente"),
        ("tecnico_tic", "0", "0", "Pendiente"),
        ("estadistico", "0", "0", "Pendiente"),
        ("jurisprudencial", "0", "0", "Pendiente"),
    ]
    for i, row in enumerate(stats, 2):
        for j, val in enumerate(row, 1):
            ws2.cell(row=i, column=j, value=val)

    ruta = CARPETA_SALIDA / "XLSX_03_Inventario_Corpus_CecilIA.xlsx"
    wb.save(str(ruta))
    return ruta


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENTOS ADICIONALES: Retroalimentacion, Informe ejecutivo, Guia pruebas
# ══════════════════════════════════════════════════════════════════════════════

def generar_formato_retroalimentacion() -> Path:
    doc = DocCGR("Formato de Retroalimentacion de Pruebas",
                 "Instrumento de evaluacion para usuarios piloto",
                 "FORMATO_Retroalimentacion_Pruebas_CecilIA")
    doc.agregar_portada()
    doc.agregar_encabezado_pagina()

    doc.h1("Datos del evaluador", "1.")
    doc.tabla(
        ["Campo", "Valor"],
        [
            ["Nombre completo", "[                                                    ]"],
            ["Cargo", "[                                                    ]"],
            ["Direccion", "[ ] DES    [ ] DVF"],
            ["Fecha de evaluacion", "[    /    /2026   ]"],
        ],
    )

    doc.h1("Evaluacion por modulo", "2.")
    modulos = ["Chat IA", "Formatos CGR", "Hallazgos", "Capacitacion", "Observatorio", "Analitica"]
    for mod in modulos:
        doc.h2(mod)
        doc.tabla(
            ["Criterio", "1", "2", "3", "4", "5"],
            [
                ["Facilidad de uso", "[ ]", "[ ]", "[ ]", "[ ]", "[ ]"],
                ["Calidad de respuestas", "[ ]", "[ ]", "[ ]", "[ ]", "[ ]"],
                ["Velocidad", "[ ]", "[ ]", "[ ]", "[ ]", "[ ]"],
                ["Utilidad para su trabajo", "[ ]", "[ ]", "[ ]", "[ ]", "[ ]"],
            ],
        )

    doc.h1("Observaciones generales", "3.")
    doc.parrafo("[Espacio para comentarios del evaluador]")
    doc.parrafo(" ")
    doc.parrafo(" ")

    doc.h1("Firma", "4.")
    doc.parrafo(" ")
    doc.parrafo("_________________________________")
    doc.parrafo("Nombre y firma del evaluador")

    doc.agregar_nota_ia()
    return doc.guardar()


def generar_informe_ejecutivo() -> Path:
    doc = DocCGR("Informe Ejecutivo de Preproduccion",
                 "Estado del piloto CecilIA v2 — Abril 2026",
                 "INFORME_Ejecutivo_Preproduccion_CecilIA")
    doc.agregar_portada()
    doc.agregar_encabezado_pagina()

    doc.h1("Resumen ejecutivo", "1.")
    doc.parrafo(
        "CecilIA v2 ha completado satisfactoriamente 11 sprints de desarrollo, resultando en un sistema "
        "funcional de inteligencia artificial para asistencia al control fiscal. El sistema incluye "
        "chat con RAG, generacion de 30 formatos CGR, gestion de hallazgos, capacitacion adaptativa, "
        "observatorio TIC y analitica. Validacion final: 17/17 endpoints OK."
    )

    doc.h1("Metricas del sistema", "2.")
    doc.tabla(
        ["Metrica", "Valor"],
        [
            ["Tablas en PostgreSQL", "16+"],
            ["Usuarios registrados", "14 (12 originales + 2 aprendices)"],
            ["Documentos en corpus", "221"],
            ["Chunks con embeddings", "10,016"],
            ["Formatos CGR implementados", "30 (15 personalizados + 15 genericos)"],
            ["Rutas de capacitacion", "4"],
            ["Endpoints API", "50+"],
            ["Sprints completados", "11"],
        ],
    )

    doc.h1("Estado por sprint", "3.")
    doc.tabla(
        ["Sprint", "Contenido", "Estado"],
        [
            ["0", "Scaffolding, auth JWT, RBAC", "Completado"],
            ["1", "Pipeline RAG", "Completado"],
            ["2", "Multi-agente LangGraph", "Completado"],
            ["3", "Frontend profesional", "Completado"],
            ["4", "Formatos CGR 1-30", "Completado"],
            ["5", "Hallazgos", "Completado"],
            ["6", "Capacitacion", "Completado"],
            ["7", "APIs externas", "Completado"],
            ["8", "Observatorio TIC", "Completado"],
            ["9-10", "Fine-tuning + Analitica", "Completado"],
            ["11", "Desktop Agent", "Completado"],
        ],
    )

    doc.h1("Proximos pasos", "4.")
    doc.lista([
        "Pruebas con usuarios reales (auditores DES y DVF)",
        "Recopilacion de retroalimentacion",
        "Ajustes basados en feedback",
        "Registro ante DNDA",
        "Despliegue en VPS de produccion",
    ])

    doc.agregar_nota_ia()
    doc.agregar_contraportada()
    return doc.guardar()


def generar_guia_pruebas() -> Path:
    doc = DocCGR("Guia de Pruebas por Rol",
                 "Instrucciones para validacion del sistema por cada rol",
                 "GUIA_Pruebas_Por_Rol_CecilIA")
    doc.agregar_portada()
    doc.agregar_encabezado_pagina()

    roles_pruebas = [
        ("Administrador (admin_tic)", [
            "Iniciar sesion con admin.cecilia / CecilIA_Admin_2026!",
            "Verificar acceso a todos los modulos",
            "Verificar que puede ver ambas direcciones (DES y DVF)",
            "Probar generacion de formatos",
            "Verificar dashboard de analitica",
        ]),
        ("Auditor DVF", [
            "Iniciar sesion con auditor.dvf.01 / Auditor_DVF_2026!",
            "Verificar que solo ve direccion DVF",
            "Hacer una consulta sobre hallazgos en el chat",
            "Generar un formato F14 (Observaciones)",
            "Crear un hallazgo con los 4 elementos",
        ]),
        ("Auditor DES", [
            "Iniciar sesion con auditor.des.01 / Auditor_DES_2026!",
            "Verificar que solo ve direccion DES",
            "Hacer una consulta sobre estudios sectoriales",
            "Verificar acceso al Observatorio TIC",
        ]),
        ("Director DVF", [
            "Iniciar sesion con jose.ramirez / DVF_Director_2026!",
            "Verificar acceso a analitica",
            "Verificar que puede aprobar hallazgos",
        ]),
        ("Aprendiz", [
            "Iniciar sesion con aprendiz.dvf / Aprendiz_DVF_2026!",
            "Verificar que solo ve modulo de Capacitacion",
            "Completar cuestionario de estilo de aprendizaje",
            "Iniciar una leccion y usar el tutor IA",
            "Generar un recurso en la Biblioteca",
        ]),
    ]

    for i, (rol, pruebas) in enumerate(roles_pruebas, 1):
        doc.h1(f"Pruebas para: {rol}", f"{i}.")
        doc.lista(pruebas, ordenada=True)
        doc.tabla(
            ["Prueba", "Resultado", "Observaciones"],
            [[f"Prueba {j}", "[ ] OK  [ ] Fallo", ""] for j in range(1, len(pruebas) + 1)],
        )

    doc.agregar_nota_ia()
    return doc.guardar()


# ══════════════════════════════════════════════════════════════════════════════
# FUNCION PRINCIPAL
# ══════════════════════════════════════════════════════════════════════════════

def main():
    print()
    print("=" * 55)
    print("  GENERACION DE DOCUMENTACION — CecilIA v2")
    print("  Contraloria General de la Republica")
    print("=" * 55)
    print()

    CARPETA_SALIDA.mkdir(parents=True, exist_ok=True)

    tareas = []

    # Doc 1-2: Manuales
    tareas.append(("01_Manual_Usuario_CecilIA_v2.docx", generar_manual_usuario))
    tareas.append(("02_Manual_Tecnico_CecilIA_v2.docx", generar_manual_tecnico))

    # Doc 3-15: Politicas
    for num, titulo, slug, desc, secciones in POLITICAS:
        tareas.append((
            f"{num}_Politica_{slug}.docx",
            lambda n=num, t=titulo, s=slug, d=desc, sec=secciones: generar_politica(n, t, s, d, sec)
        ))

    # Doc 16-17
    tareas.append(("16_Registro_Software_DNDA_CecilIA.docx", generar_registro_dnda))
    tareas.append(("17_Guia_Implementacion_Institucional_CecilIA.docx", generar_guia_implementacion))

    # XLSX
    tareas.append(("XLSX_01_Matriz_RACI_CecilIA.xlsx", generar_xlsx_raci))
    tareas.append(("XLSX_02_Inventario_API_CecilIA.xlsx", generar_xlsx_api))
    tareas.append(("XLSX_03_Inventario_Corpus_CecilIA.xlsx", generar_xlsx_corpus))

    # Documentos adicionales
    tareas.append(("FORMATO_Retroalimentacion_Pruebas_CecilIA.docx", generar_formato_retroalimentacion))
    tareas.append(("INFORME_Ejecutivo_Preproduccion_CecilIA.docx", generar_informe_ejecutivo))
    tareas.append(("GUIA_Pruebas_Por_Rol_CecilIA.docx", generar_guia_pruebas))

    total = len(tareas)
    exitosos = 0
    resultados = []

    for i, (nombre, func) in enumerate(tareas, 1):
        print(f"  [{i:2d}/{total}] Generando: {nombre}...", end=" ", flush=True)
        inicio = time.time()
        try:
            ruta = func()
            duracion = time.time() - inicio
            tamano = ruta.stat().st_size if ruta.exists() else 0
            tamano_kb = tamano / 1024
            print(f"OK ({tamano_kb:.0f} KB, {duracion:.1f}s)")
            exitosos += 1
            resultados.append((nombre, tamano_kb, True))
        except Exception as e:
            print(f"ERROR: {e}")
            resultados.append((nombre, 0, False))

    print()
    print("=" * 55)
    print(f"  RESUMEN: {exitosos}/{total} documentos generados")
    print(f"  17 DOCX + 3 XLSX + 3 documentos adicionales")
    print(f"  Ubicacion: {CARPETA_SALIDA}")
    print("=" * 55)
    print()
    print("  Archivos generados:")
    for nombre, tamano, ok in resultados:
        estado = "OK" if ok else "ERROR"
        print(f"    [{estado}] {nombre} ({tamano:.0f} KB)")
    print()


if __name__ == "__main__":
    main()
